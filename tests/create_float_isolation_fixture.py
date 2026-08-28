"""Generate spec fixture: H1 Intro + floats A-E as per task description."""
import os, io, zipfile, xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm
from PIL import Image as PILImage

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIX = os.path.join(PROJ, "tests", "fixtures")
TMP = os.path.join(PROJ, "tests", ".flt_iso_tmp")
os.makedirs(FIX, exist_ok=True)
os.makedirs(TMP, exist_ok=True)

WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

def _pil_bytes(size, color):
    img = PILImage.new("RGB", size, color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()

def _img_file(name, data):
    p = os.path.join(TMP, name)
    with open(p, "wb") as f:
        f.write(data)
    return p

RED = _pil_bytes((200,150),(200,30,30))
BLUE = _pil_bytes((200,150),(30,60,200))
GREEN = _pil_bytes((200,150),(30,160,60))
YELLOW = _pil_bytes((200,150),(220,180,20))
PURP = _pil_bytes((200,150),(120,40,160))

RED_F = _img_file("iso_red.png", RED)
BLUE_F = _img_file("iso_blue.png", BLUE)
GREEN_F = _img_file("iso_green.png", GREEN)
YELLOW_F = _img_file("iso_yellow.png", YELLOW)
PURP_F = _img_file("iso_purp.png", PURP)

def _convert_inline_to_anchor(docx_path, specs):
    with zipfile.ZipFile(docx_path, "r") as z:
        stored = {i.filename: z.read(i.filename) for i in z.infolist()}
    root = ET.fromstring(stored["word/document.xml"])
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    WPq = "{%s}" % WP
    inlines = list(root.iter(WPq + "inline"))
    assert len(inlines) == len(specs), "expected %d inlines got %d" % (len(specs), len(inlines))
    for inline, spec in zip(inlines, specs):
        anchor = ET.Element(WPq + "anchor", {
            "distT":"0","distB":"0","distL":"0","distR":"0",
            "simplePos":"0","relativeHeight":"251658240",
            "behindDoc":"0","locked":"0","layoutInCell":"1","allowOverlap":"1",
        })
        for child in list(inline):
            anchor.append(child)
        sp = ET.SubElement(anchor, WPq + "simplePos")
        sp.set("x","0"); sp.set("y","0")
        ph = ET.SubElement(anchor, WPq + "positionH")
        ph.set("relativeFrom", spec["relH"])
        if spec.get("alignH"):
            a = ET.SubElement(ph, WPq + "align"); a.text = spec["alignH"]
        else:
            o = ET.SubElement(ph, WPq + "posOffset"); o.text = str(spec.get("offH",0))
        pv = ET.SubElement(anchor, WPq + "positionV")
        pv.set("relativeFrom", spec["relV"])
        if spec.get("alignV"):
            a = ET.SubElement(pv, WPq + "align"); a.text = spec["alignV"]
        else:
            o = ET.SubElement(pv, WPq + "posOffset"); o.text = str(spec.get("offV",0))
        wrap = ET.SubElement(anchor, WPq + "wrapNone")
        parent = None
        for d in root.iter(W + "drawing"):
            if inline in list(d):
                parent = d
                break
        idx = list(parent).index(inline)
        parent.remove(inline)
        parent.insert(idx, anchor)
    stored["word/document.xml"] = ET.tostring(root, encoding="unicode").encode("utf-8")
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as z2:
        for fname, data in stored.items():
            z2.writestr(fname, data)

doc = Document()
# H1 Introduction
doc.add_heading("Introduction", level=1)
p = doc.add_paragraph("Introduction paragraph belonging to H1 Introduction.")
# floating image A right after introduction paragraph (owning paragraph is the next paragraph that holds the drawing)
pa = doc.add_paragraph()
pa.add_run().add_picture(RED_F, width=Cm(4), height=Cm(3))
pa.add_run(" Image A anchor text")

# H2 Architecture
doc.add_heading("Architecture", level=2)
p = doc.add_paragraph("Architecture paragraph under H2 Architecture.")
pb = doc.add_paragraph()
pb.add_run().add_picture(BLUE_F, width=Cm(4), height=Cm(3))
pb.add_run(" Image B anchor text")

# H3 Data Model
doc.add_heading("Data Model", level=3)
p = doc.add_paragraph("Data model paragraph under H3 Data Model.")
pc = doc.add_paragraph()
pc.add_run().add_picture(GREEN_F, width=Cm(4), height=Cm(3))
pc.add_run(" Image C anchor text")

# H2 Implementation
doc.add_heading("Implementation", level=2)
p = doc.add_paragraph("Implementation paragraph under H2 Implementation.")
pd = doc.add_paragraph()
pd.add_run().add_picture(YELLOW_F, width=Cm(4), height=Cm(3))
pd.add_run(" Image D anchor text")

# H1 Conclusion
doc.add_heading("Conclusion", level=1)
p = doc.add_paragraph("Conclusion paragraph under H1 Conclusion.")
pe = doc.add_paragraph()
pe.add_run().add_picture(PURP_F, width=Cm(4), height=Cm(3))
pe.add_run(" Image E anchor text")

out = os.path.join(FIX, "float-isolation.docx")
doc.save(out)
# All 5 images become page-relative anchors with distinct offsets
specs = [
    {"relH":"page","offH":800000,"relV":"page","offV":300000},
    {"relH":"page","offH":1500000,"relV":"page","offV":600000},
    {"relH":"page","offH":2200000,"relV":"page","offV":900000},
    {"relH":"page","offH":3000000,"relV":"page","offV":1200000},
    {"relH":"page","offH":3800000,"relV":"page","offV":1500000},
]
_convert_inline_to_anchor(out, specs)
print("wrote", out)
