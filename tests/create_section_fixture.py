"""Generate a real DOCX fixture for the isolated-heading-view browser test.

Structure required by the feature spec:
  H1 Introduction
    paragraph
  H2 Architecture
    paragraph
    table
    hyperlink
    inline image
  H3 Data Model
    paragraph
  H2 Implementation
    paragraph
  H1 Conclusion
    paragraph
  H2 Summary
    paragraph
"""
import os
from docx import Document
from docx.shared import Pt

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIX = os.path.join(PROJ, "tests", "fixtures")
os.makedirs(FIX, exist_ok=True)

# Tiny valid red PNG for inline-image content fidelity check.
from PIL import Image
PNG_PATH = os.path.join(FIX, "tiny.png")
Image.new("RGB", (40, 40), (200, 40, 40)).save(PNG_PATH)


def add_hyperlink(paragraph, url, text):
    # python-docx has no direct add_hyperlink; build the run + relationship.
    from docx.oxml.shared import qn
    from docx.oxml.ns import nsmap  # noqa
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    new_run = paragraph.add_run()
    new_run.text = text
    hyperlink.append(new_run._r)
    paragraph._p.append(hyperlink)
    return hyperlink


doc = Document()

# H1 Introduction
doc.add_heading("Introduction", level=1)
doc.add_paragraph("Intro body text that belongs to the Introduction section.")

# H2 Architecture
doc.add_heading("Architecture", level=2)
doc.add_paragraph("Architecture overview paragraph under the Architecture heading.")
# table inside Architecture
tbl = doc.add_table(rows=2, cols=2)
tbl.style = "Table Grid"
tbl.cell(0, 0).text = "Arch A1"
tbl.cell(0, 1).text = "Arch B1"
tbl.cell(1, 0).text = "Arch A2"
tbl.cell(1, 1).text = "Arch B2"
# hyperlink inside Architecture
p = doc.add_paragraph()
add_hyperlink(p, "https://example.com", "Example hyperlink in Architecture")
# inline image inside Architecture
doc.add_picture(PNG_PATH, width=Pt(40))

# H3 Data Model
doc.add_heading("Data Model", level=3)
doc.add_paragraph("Data model details under the Data Model heading.")

# H2 Implementation
doc.add_heading("Implementation", level=2)
doc.add_paragraph("Implementation notes under the Implementation heading.")

# H1 Conclusion
doc.add_heading("Conclusion", level=1)
doc.add_paragraph("Conclusion text under the Conclusion heading.")

# H2 Summary
doc.add_heading("Summary", level=2)
doc.add_paragraph("Summary text under the Summary heading.")

out = os.path.join(FIX, "section-isolation.docx")
doc.save(out)
print("wrote", out)
