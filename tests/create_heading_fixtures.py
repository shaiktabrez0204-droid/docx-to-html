"""Generate REAL .docx fixtures for the styles/heading/hierarchy/TOC pipeline.

All files are valid OOXML produced by python-docx 1.2.0.
"""
import os
import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

BASE = r"C:\Tabrez\docx-html\docx-to-html\tests\fixtures"


def save(doc, name):
    path = os.path.join(BASE, name)
    doc.save(path)
    print("OK  ", name)
    return path


def make_h1_only():
    d = Document()
    d.add_heading("Introduction", level=1)
    d.add_paragraph("Body text under intro.")
    return d


def make_h1_h2():
    d = Document()
    d.add_heading("Chapter One", level=1)
    d.add_paragraph("Some text.")
    d.add_heading("Section A", level=2)
    d.add_paragraph("More text.")
    d.add_heading("Section B", level=2)
    d.add_paragraph("End.")
    return d


def make_h1_h2_h3():
    d = Document()
    d.add_heading("Top", level=1)
    d.add_heading("Mid", level=2)
    d.add_heading("Bottom", level=3)
    d.add_paragraph("Detail.")
    return d


def make_h1_h6():
    d = Document()
    for lvl, txt in [(1, "One"), (2, "Two"), (3, "Three"),
                    (4, "Four"), (5, "Five"), (6, "Six")]:
        d.add_heading(txt, level=lvl)
    return d


def make_numbered():
    d = Document()
    d.add_heading("1. Getting Started", level=1)
    d.add_heading("1.1 Setup", level=2)
    d.add_heading("1.1.1 Install", level=3)
    d.add_heading("1.2 Configuration", level=2)
    return d


def make_unnumbered():
    d = Document()
    d.add_heading("Overview", level=1)
    d.add_heading("Components", level=2)
    d.add_heading("Database", level=3)
    return d


def make_skipped():
    d = Document()
    d.add_heading("Part One", level=1)
    d.add_heading("Deep Detail", level=3)   # skip level 2
    d.add_heading("Section Two", level=2)   # H2 after H3
    return d


def make_duplicate():
    d = Document()
    d.add_heading("Introduction", level=1)
    d.add_paragraph("x")
    d.add_heading("Summary", level=2)
    d.add_paragraph("y")
    d.add_heading("Introduction", level=1)  # duplicate text, H1
    d.add_paragraph("z")
    d.add_heading("Summary", level=2)       # duplicate text, H2
    return d


def make_custom():
    d = Document()
    style = d.styles.add_style("My Custom Heading", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = d.styles["Heading 1"]  # inherits Heading 1 outline level via BasedOn chain
    d.add_paragraph("Custom Heading Text", style=style)
    d.add_paragraph("Body after custom.")
    d.add_paragraph("Another Custom Heading", style=style)
    # report resolved styleId
    for s in d.styles:
        if s.name == "My Custom Heading":
            print("INFO custom styleId =", repr(s.style_id))
    return d


def make_mixed():
    d = Document()
    # Realistic Word auto-TOC uses style "TOC 1"; simulate it. Must NOT become a heading.
    if "TOC 1" not in [s.name for s in d.styles]:
        toc_style = d.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
        toc_style.base_style = d.styles["Normal"]
    d.add_heading("Introduction", level=1)
    d.add_paragraph("The intro paragraph.")
    d.add_heading("Architecture", level=2)
    d.add_paragraph("Architecture body.")
    d.add_heading("Data Model", level=3)
    d.add_paragraph("Data model body.")
    d.add_heading("Implementation", level=2)
    d.add_paragraph("Implementation body.")
    # Simulated Word TOC artifact — must NOT become a heading
    d.add_paragraph("Table of Contents Entry............ 3", style="TOC 1")
    # Visually prominent NON-heading — must NOT become a heading
    p = d.add_paragraph()
    run = p.add_run("This is a big bold statement.")
    run.bold = True
    run.font.size = docx.shared.Pt(28)
    d.add_heading("Conclusion", level=1)
    d.add_paragraph("The end.")
    return d


if __name__ == "__main__":
    save(make_h1_only(), "h1-only.docx")
    save(make_h1_h2(), "h1-h2.docx")
    save(make_h1_h2_h3(), "h1-h2-h3.docx")
    save(make_h1_h6(), "h1-h6.docx")
    save(make_numbered(), "numbered-headings.docx")
    save(make_unnumbered(), "unnumbered-headings.docx")
    save(make_skipped(), "skipped-levels.docx")
    save(make_duplicate(), "duplicate-heading-text.docx")
    save(make_custom(), "custom-heading-style.docx")
    save(make_mixed(), "mixed-document.docx")
    print("ALL FIXTURES WRITTEN")
