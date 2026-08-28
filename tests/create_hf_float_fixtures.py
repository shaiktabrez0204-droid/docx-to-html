import os, io, zipfile, xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm
from PIL import Image as PILImage

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIX = os.path.join(PROJ, "tests", "fixtures")
TMP = os.path.join(PROJ, "tests", ".hf_tmp")
os.makedirs(FIX, exist_ok=True)
os.makedirs(TMP, exist_ok=True)
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

def pil_bytes(size, color):
    im = PILImage.new("RGB", size, color)
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return buf.getvalue()

def img_file(name, data):
    p=os.path.join(TMP, name)
    with open(p,"wb") as f: f.write(data)
    return p

RED = pil_bytes((200,150),(200,30,30))
BLUE = pil_bytes((200,150),(30,60,200))
GREEN = pil_bytes((200,150),(30,160,60))
YELL = pil_bytes((200,150),(220,180,20))
PURP = pil_bytes((200,150),(120,40,160))
CYAN = pil_bytes((200,150),(20,160,160))

files = {k: img_file(k+".png", v) for k,v in [("red",RED),("blue",BLUE),("green",GREEN),("yell",YELL),("purp",PURP),("cyan",CYAN)]}

def convert_part(docx_path, part_name, specs):
    with zipfile.ZipFile(docx_path,"r") as z:
        stored={i.filename: z.read(i.filename) for i in z.infolist()}
    if part_name not in stored:
        print(f"WARN {part_name} not in doc")
        return
    root=ET.fromstring(stored[part_name])
    W="{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    WPq="{%s}"%WP
    inlines=list(root.iter(WPq+"inline"))
    # specs length must match inlines in that part
    assert len(inlines)==len(specs), f"{part_name}: expected {len(specs)} inlines got {len(inlines)}"
    for inline, spec in zip(inlines, specs):
        anchor=ET.Element(WPq+"anchor", {"distT":"0","distB":"0","distL":"0","distR":"0","simplePos":"0","relativeHeight":"251658240","behindDoc":"0","locked":"0","layoutInCell":"1","allowOverlap":"1"})
        for child in list(inline):
            anchor.append(child)
        sp=ET.SubElement(anchor, WPq+"simplePos"); sp.set("x","0"); sp.set("y","0")
        ph=ET.SubElement(anchor, WPq+"positionH"); ph.set("relativeFrom", spec["relH"])
        o=ET.SubElement(ph, WPq+"posOffset"); o.text=str(spec.get("offH",0))
        pv=ET.SubElement(anchor, WPq+"positionV"); pv.set("relativeFrom", spec["relV"])
        o2=ET.SubElement(pv, WPq+"posOffset"); o2.text=str(spec.get("offV",0))
        wrap=ET.SubElement(anchor, WPq+"wrapNone")
        parent=None
        for d in root.iter(W+"drawing"):
            if inline in list(d):
                parent=d; break
        idx=list(parent).index(inline)
        parent.remove(inline); parent.insert(idx, anchor)
    stored[part_name]=ET.tostring(root, encoding="unicode").encode("utf-8")
    with zipfile.ZipFile(docx_path,"w", zipfile.ZIP_DEFLATED) as z2:
        for fn, data in stored.items():
            z2.writestr(fn, data)

# CASE A: simple doc with header/footer floats + body sections with table floats
def build_case_a():
    doc=Document()
    sec=doc.sections[0]
    # header
    hdr=sec.header
    hdr.is_linked_to_previous=False
    p=hdr.paragraphs[0]
    r=p.add_run()
    r.add_picture(files["cyan"], width=Cm(3), height=Cm(2))
    p.add_run(" Header float")
    # footer
    ftr=sec.footer
    ftr.is_linked_to_previous=False
    p2=ftr.paragraphs[0]
    r2=p2.add_run()
    r2.add_picture(files["purp"], width=Cm(3), height=Cm(2))
    p2.add_run(" Footer float")
    # body
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Intro para")
    tbl=doc.add_table(rows=1, cols=1); tbl.style="Table Grid"
    cell=tbl.cell(0,0); cell.text=""
    p=cell.paragraphs[0]; r=p.add_run(); r.add_picture(files["red"], width=Cm(4), height=Cm(3)); p.add_run(" A")
    doc.add_heading("Architecture", level=2)
    doc.add_paragraph("Arch para")
    tbl=doc.add_table(rows=1, cols=1); tbl.style="Table Grid"
    cell=tbl.cell(0,0); cell.text=""
    p=cell.paragraphs[0]; r=p.add_run(); r.add_picture(files["blue"], width=Cm(4), height=Cm(3)); p.add_run(" B")
    doc.add_heading("Implementation", level=2)
    doc.add_paragraph("Impl para")
    tbl=doc.add_table(rows=1, cols=1); tbl.style="Table Grid"
    cell=tbl.cell(0,0); cell.text=""
    p=cell.paragraphs[0]; r=p.add_run(); r.add_picture(files["green"], width=Cm(4), height=Cm(3)); p.add_run(" C")
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("Concl para")
    tbl=doc.add_table(rows=1, cols=1); tbl.style="Table Grid"
    cell=tbl.cell(0,0); cell.text=""
    p=cell.paragraphs[0]; r=p.add_run(); r.add_picture(files["yell"], width=Cm(4), height=Cm(3)); p.add_run(" D")
    out=os.path.join(FIX, "hf-float-caseA.docx")
    doc.save(out)
    # convert body inlines (4) and header/footer inlines (1 each)
    # Body has 4 images (tables) -> document.xml
    convert_part(out, "word/document.xml", [
        {"relH":"page","offH":800000,"relV":"page","offV":300000},
        {"relH":"page","offH":1500000,"relV":"page","offV":600000},
        {"relH":"page","offH":2200000,"relV":"page","offV":900000},
        {"relH":"page","offH":3800000,"relV":"page","offV":1200000},
    ])
    convert_part(out, "word/header1.xml", [{"relH":"page","offH":500000,"relV":"page","offV":200000}])
    convert_part(out, "word/footer1.xml", [{"relH":"page","offH":500000,"relV":"page","offV":500000}])
    print("wrote", out)
    return out

# CASE B: multiple sections with distinct headers
def build_case_b():
    doc=Document()
    # Sec 0 header
    sec0=doc.sections[0]
    hdr0=sec0.header
    hdr0.is_linked_to_previous=False
    p=hdr0.paragraphs[0]
    r=p.add_run()
    r.add_picture(files["red"], width=Cm(3), height=Cm(2))
    p.add_run(" Header Sec0")
    ftr0=sec0.footer
    ftr0.is_linked_to_previous=False
    p=ftr0.paragraphs[0]
    r=p.add_run(); r.add_picture(files["blue"], width=Cm(3), height=Cm(2)); p.add_run(" Footer Sec0")

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Intro sec0")
    tbl=doc.add_table(rows=1, cols=1); tbl.style="Table Grid"
    cell=tbl.cell(0,0); cell.text=""; p=cell.paragraphs[0]; r=p.add_run(); r.add_picture(files["cyan"], width=Cm(3), height=Cm(2)); p.add_run(" body A sec0")

    # New section
    sec1=doc.add_section()
    sec1.header.is_linked_to_previous=False
    sec1.footer.is_linked_to_previous=False
    hdr1=sec1.header
    p=hdr1.paragraphs[0]
    r=p.add_run(); r.add_picture(files["green"], width=Cm(3), height=Cm(2)); p.add_run(" Header Sec1")
    ftr1=sec1.footer
    p=ftr1.paragraphs[0]
    r=p.add_run(); r.add_picture(files["yell"], width=Cm(3), height=Cm(2)); p.add_run(" Footer Sec1")

    doc.add_heading("Architecture", level=2)
    doc.add_paragraph("Arch sec1")
    tbl=doc.add_table(rows=1, cols=1); tbl.style="Table Grid"
    cell=tbl.cell(0,0); cell.text=""; p=cell.paragraphs[0]; r=p.add_run(); r.add_picture(files["purp"], width=Cm(3), height=Cm(2)); p.add_run(" body B sec1")

    out=os.path.join(FIX, "hf-float-caseB.docx")
    doc.save(out)
    # document has 2 body images
    convert_part(out, "word/document.xml", [
        {"relH":"page","offH":800000,"relV":"page","offV":300000},
        {"relH":"page","offH":1500000,"relV":"page","offV":600000},
    ])
    convert_part(out, "word/header1.xml", [{"relH":"page","offH":400000,"relV":"page","offV":150000}])
    convert_part(out, "word/header2.xml", [{"relH":"page","offH":400000,"relV":"page","offV":150000}])
    convert_part(out, "word/footer1.xml", [{"relH":"page","offH":400000,"relV":"page","offV":400000}])
    convert_part(out, "word/footer2.xml", [{"relH":"page","offH":400000,"relV":"page","offV":400000}])
    print("wrote", out)
    return out

if __name__=="__main__":
    build_case_a()
    build_case_b()
