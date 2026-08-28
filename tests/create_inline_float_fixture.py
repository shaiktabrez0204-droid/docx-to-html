import os, io, zipfile, xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm
from PIL import Image as PILImage

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIX = os.path.join(PROJ, "tests", "fixtures")
TMP = os.path.join(PROJ, "tests", ".inline_tmp")
os.makedirs(FIX, exist_ok=True)
os.makedirs(TMP, exist_ok=True)
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

def pil_bytes(size, color):
    im = PILImage.new("RGB", size, color)
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return buf.getvalue()

def img_file(name, data):
    p = os.path.join(TMP, name)
    with open(p, "wb") as f:
        f.write(data)
    return p

RED = pil_bytes((200,150),(200,30,30))
BLUE = pil_bytes((200,150),(30,60,200))
GREEN = pil_bytes((200,150),(30,160,60))
YELL = pil_bytes((200,150),(220,180,20))
PURP = pil_bytes((200,150),(120,40,160))

files = {k: img_file(k+".png", v) for k,v in [("red",RED),("blue",BLUE),("green",GREEN),("yell",YELL),("purp",PURP)]}

def convert_to_wrap_anchor(docx_path, specs):
    with zipfile.ZipFile(docx_path, "r") as z:
        stored = {i.filename: z.read(i.filename) for i in z.infolist()}
    root = ET.fromstring(stored["word/document.xml"])
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    WPq = "{%s}" % WP
    inlines = list(root.iter(WPq + "inline"))
    assert len(inlines) == len(specs), "expected %d got %d" % (len(specs), len(inlines))
    for inline, spec in zip(inlines, specs):
        anchor = ET.Element(WPq + "anchor", {
            "distT":"0","distB":"0","distL":"0","distR":"0",
            "simplePos":"0","relativeHeight":"251658240",
            "behindDoc":"0","locked":"0","layoutInCell":"1","allowOverlap":"1",
        })
        for child in list(inline):
            anchor.append(child)
        sp = ET.SubElement(anchor, WPq + "simplePos"); sp.set("x","0"); sp.set("y","0")
        ph = ET.SubElement(anchor, WPq + "positionH"); ph.set("relativeFrom", spec["relH"])
        align = ET.SubElement(ph, WPq + "align"); align.text = spec["alignH"]
        pv = ET.SubElement(anchor, WPq + "positionV"); pv.set("relativeFrom", spec["relV"])
        alignV = ET.SubElement(pv, WPq + "align"); alignV.text = spec.get("alignV","top")
        # wrapSquare with bothSides
        wrap = ET.SubElement(anchor, WPq + "wrapSquare")
        wrap.set("wrapText", "bothSides")
        wrap.set("distT", "0"); wrap.set("distB", "0"); wrap.set("distL", "0"); wrap.set("distR", "0")
        ET.SubElement(wrap, WPq + "effectExtent", {"l":"0","t":"0","r":"0","b":"0"})
        parent = None
        for d in root.iter(W + "drawing"):
            if inline in list(d):
                parent = d
                break
        idx = list(parent).index(inline)
        parent.remove(inline)
        parent.insert(idx, anchor)
    stored["word/document.xml"] = ET.tostring(root, encoding="unicode").encode("utf-8")
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as z2:
        for fn, data in stored.items():
            z2.writestr(fn, data)

doc = Document()
doc.add_heading("Introduction", level=1)
p = doc.add_paragraph("Introduction paragraph with wrapSquare float-left A. " + "Lorem ipsum dolor sit amet consectetur adipiscing elit. " * 6)
pa = doc.add_paragraph()
pa.add_run().add_picture(files["red"], width=Cm(4), height=Cm(3))
pa.add_run(" Image A float-left text surrounding to show wrapping. " + "More text to wrap around the image. "*5)

doc.add_heading("Architecture", level=2)
p = doc.add_paragraph("Architecture paragraph with wrapSquare float-right B. " + "Lorem ipsum dolor sit amet. "*6)
pb = doc.add_paragraph()
pb.add_run().add_picture(files["blue"], width=Cm(4), height=Cm(3))
pb.add_run(" Image B float-right " + "Wrapping text content. "*5)

doc.add_heading("Data Model", level=3)
p = doc.add_paragraph("Data model paragraph with wrapSquare float-left C. " + "Lorem ipsum. "*6)
pc = doc.add_paragraph()
pc.add_run().add_picture(files["green"], width=Cm(4), height=Cm(3))
pc.add_run(" Image C float-left " + "Extra wrapping text. "*5)

doc.add_heading("Implementation", level=2)
p = doc.add_paragraph("Implementation paragraph with wrapSquare float-right D. " + "Lorem ipsum. "*6)
pd = doc.add_paragraph()
pd.add_run().add_picture(files["yell"], width=Cm(4), height=Cm(3))
pd.add_run(" Image D float-right " + "More wrapping text. "*5)

doc.add_heading("Conclusion", level=1)
p = doc.add_paragraph("Conclusion paragraph with wrapSquare float-left E. " + "Lorem ipsum. "*6)
pe = doc.add_paragraph()
pe.add_run().add_picture(files["purp"], width=Cm(4), height=Cm(3))
pe.add_run(" Image E float-left " + "Final wrapping text. "*5)

out = os.path.join(FIX, "inline-float-isolation.docx")
doc.save(out)
specs = [
    {"relH":"page","alignH":"left","relV":"page","alignV":"top"},
    {"relH":"page","alignH":"right","relV":"page","alignV":"top"},
    {"relH":"page","alignH":"left","relV":"page","alignV":"top"},
    {"relH":"page","alignH":"right","relV":"page","alignV":"top"},
    {"relH":"page","alignH":"left","relV":"page","alignV":"top"},
]
convert_to_wrap_anchor(out, specs)
print("wrote", out)
