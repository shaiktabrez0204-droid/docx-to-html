"""Create REAL DOCX fixture with floating images INSIDE comment ranges."""
import os, io, zipfile, xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm
from PIL import Image as PILImage

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIX = os.path.join(PROJECT_ROOT, "tests", "fixtures")
TMP = os.path.join(PROJECT_ROOT, "tests", ".cflt_tmp")
os.makedirs(FIX, exist_ok=True)
os.makedirs(TMP, exist_ok=True)

WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

def _pil_bytes(size, color):
    img = PILImage.new("RGB", size, color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()

def _img_file(name, data):
    p = os.path.join(TMP, name)
    with open(p, "wb") as f:
        f.write(data)
    return p

RED = _pil_bytes((300,200),(200,30,30))
BLUE = _pil_bytes((160,120),(30,60,200))
GREEN = _pil_bytes((100,100),(30,160,60))
YELLOW = _pil_bytes((260,180),(220,180,20))
PURP = _pil_bytes((200,150),(120,40,160))
CYAN = _pil_bytes((240,160),(20,160,160))
ORANGE = _pil_bytes((280,200),(220,120,20))

RED_F = _img_file("cfr.png", RED)
BLUE_F = _img_file("cfb.png", BLUE)
GREEN_F = _img_file("cfg.png", GREEN)
YELLOW_F = _img_file("cfy.png", YELLOW)
PURP_F = _img_file("cfp.png", PURP)
CYAN_F = _img_file("cfc.png", CYAN)
ORANGE_F = _img_file("cfo.png", ORANGE)

def _convert_inline_to_anchor(docx_path, specs):
    with zipfile.ZipFile(docx_path, "r") as z:
        stored = {i.filename: z.read(i.filename) for i in z.infolist()}
    root = ET.fromstring(stored["word/document.xml"])
    W = "{%s}" % W_NS
    WPq = "{%s}" % WP
    inlines = list(root.iter(WPq + "inline"))
    assert len(inlines) == len(specs), f"expected {len(specs)} inlines got {len(inlines)}"
    for inline, spec in zip(inlines, specs):
        anchor = ET.Element(WPq + "anchor", {
            "distT": str(spec.get("distT", 0)),
            "distB": str(spec.get("distB", 0)),
            "distL": str(spec.get("distL", 0)),
            "distR": str(spec.get("distR", 0)),
            "simplePos": "0",
            "relativeHeight": spec.get("relativeHeight", "251658240"),
            "behindDoc": "1" if spec.get("behind") else "0",
            "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
        })
        for child in list(inline):
            anchor.append(child)
        sp = ET.SubElement(anchor, WPq + "simplePos"); sp.set("x","0"); sp.set("y","0")
        if "relH" in spec:
            ph = ET.SubElement(anchor, WPq + "positionH"); ph.set("relativeFrom", spec["relH"])
            if spec.get("alignH"):
                a = ET.SubElement(ph, WPq + "align"); a.text = spec["alignH"]
            else:
                o = ET.SubElement(ph, WPq + "posOffset"); o.text = str(spec.get("offH",0))
        if "relV" in spec:
            pv = ET.SubElement(anchor, WPq + "positionV"); pv.set("relativeFrom", spec["relV"])
            if spec.get("alignV"):
                a = ET.SubElement(pv, WPq + "align"); a.text = spec["alignV"]
            else:
                o = ET.SubElement(pv, WPq + "posOffset"); o.text = str(spec.get("offV",0))
        wk = spec.get("wrap", "Square")
        wrap = ET.SubElement(anchor, WPq + "wrap" + wk)
        if wk == "Square":
            wrap.set("wrapText", spec.get("wrapText","bothSides"))
        # polygon for Tight/Through
        poly = spec.get("polygon")
        if poly:
            wpEl = ET.SubElement(wrap, WPq + "wrapPolygon")
            wpEl.set("edited","0")
            start = ET.SubElement(wpEl, WPq + "start"); start.set("x", str(poly[0][0])); start.set("y", str(poly[0][1]))
            for x,y in poly[1:]:
                lt = ET.SubElement(wpEl, WPq + "lineTo"); lt.set("x", str(x)); lt.set("y", str(y))
        parent = None
        for d in root.iter(W + "drawing"):
            if inline in list(d):
                parent = d; break
        idx = list(parent).index(inline)
        parent.remove(inline); parent.insert(idx, anchor)
    stored["word/document.xml"] = ET.tostring(root, encoding="unicode").encode("utf-8")
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as z2:
        for fname, data in stored.items():
            z2.writestr(fname, data)

doc = Document()
# Set narrow margins so page-relative geometry is clear but not strictly needed
sec = doc.sections[0]
sec.page_width = Cm(21.59); sec.page_height = Cm(27.94)
sec.left_margin = Cm(2); sec.right_margin = Cm(2); sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)

doc.add_heading("Comment Float Test", level=1)
doc.add_paragraph("Intro before comment ranges.")

# Create 7 paragraphs each with one image
# 0..5 will be inside comment ranges, 6 is control outside
labels = ["Square page offset", "Square margin left-float", "Tight polygon", "Through polygon", "TopAndBottom paragraph", "None overlay page", "Control wrapNone outside"]
files = [RED_F, BLUE_F, GREEN_F, YELLOW_F, PURP_F, CYAN_F, ORANGE_F]
for lbl, fpath in zip(labels, files):
    p = doc.add_paragraph()
    p.add_run(lbl + " text before. ")
    p.add_run().add_picture(fpath, width=Cm(4), height=Cm(3))
    p.add_run(" text after " + lbl.lower() + ".")

out = os.path.join(FIX, "comment-float.docx")
doc.save(out)

# Define anchor specs matching 7 inlines in order
specs = [
    # 0: Square, page offset explicit, dist non-zero, explicit extent (via Cm 4x3)
    {"relH":"page","offH": 1500000, "relV":"page","offV": 800000, "wrap":"Square", "distT":114300,"distB":114300,"distL":114300,"distR":114300},
    # 1: Square with margin left align -> real float
    {"relH":"margin","alignH":"left", "relV":"margin","offV": 200000, "wrap":"Square", "distT":90000,"distB":90000,"distL":90000,"distR":90000},
    # 2: Tight with polygon, page offset
    {"relH":"page","offH": 2500000, "relV":"page","offV": 1200000, "wrap":"Tight", "distT":50000,"distB":50000,"distL":50000,"distR":50000,
     "polygon": [(0,0),(21600,0),(21600,15000),(10800,21600),(0,15000),(0,0)]},
    # 3: Through with polygon, margin offset
    {"relH":"margin","offH": 600000, "relV":"margin","offV": 400000, "wrap":"Through", "distT":60000,"distB":60000,"distL":60000,"distR":60000,
     "polygon": [(5400,0),(16200,0),(21600,10800),(16200,21600),(5400,21600),(0,10800),(5400,0)]},
    # 4: TopAndBottom, paragraph relative offset
    {"relH":"paragraph","offH": 300000, "relV":"paragraph","offV": 100000, "wrap":"TopAndBottom", "distT":80000,"distB":80000,"distL":0,"distR":0},
    # 5: None overlay, page offset, behindDoc
    {"relH":"page","offH": 3200000, "relV":"page","offV": 2000000, "wrap":"None", "distT":0,"distB":0,"distL":0,"distR":0, "behind":"1"},
    # 6: Control Outside: wrapNone page offset
    {"relH":"page","offH": 1000000, "relV":"page","offV": 500000, "wrap":"None", "distT":0,"distB":0,"distL":0,"distR":0},
]

_convert_inline_to_anchor(out, specs)
print("Anchors converted", out)

# Now inject comment ranges and comments.xml
with zipfile.ZipFile(out, "r") as z:
    stored = {i.filename: z.read(i.filename) for i in z.infolist()}

root = ET.fromstring(stored["word/document.xml"])
W = "{%s}" % W_NS
WPq = "{%s}" % WP
# Find body paragraphs in order
body = root.find(W + "body")
paras = [c for c in body if c.tag == W + "p"]
# paras[0]=heading, paras[1]=intro, paras[2..8] = 7 image paras (2-> image0, 3->image1, 4->2, 5->3, 6->4, 7->5, 8->6 control)
# Map comment ids 10..15 to image indices 0..5
comment_ids = ["10","11","12","13","14","15"]
for idx, cid in enumerate(comment_ids):
    p_elem = paras[2+idx]  # 0-based offset
    # Find w:drawing parent and its w:r container
    # Structure: w:p -> w:r -> w:drawing -> wp:anchor ...  We will wrap with commentRangeStart/End as direct w:p children OR w:r nested.
    # Use direct w:p children: insert commentRangeStart before first w:r, commentRangeEnd after last w:r, and add commentReference run after.
    # Insert Start
    start = ET.Element(W + "commentRangeStart"); start.set(W + "id", cid)
    end = ET.Element(W + "commentRangeEnd"); end.set(W + "id", cid)
    # Find index of first w:r and last w:r
    children = list(p_elem)
    first_r_idx = next((i for i,c in enumerate(children) if c.tag == W + "r"), None)
    # Insert start before first r but after pPr if present
    # Keep pPr at 0 if exists
    insert_idx = 0
    if children and children[0].tag == W + "pPr":
        insert_idx = 1
        if first_r_idx is not None and first_r_idx > insert_idx:
            insert_idx = first_r_idx
        p_elem.insert(insert_idx, start)
    else:
        p_elem.insert(insert_idx, start)
    # For end: append before commentReference, after all r's
    # Need to find updated children
    children2 = list(p_elem)
    # Insert end before we add reference, at end of r's
    # Find last r index
    last_r_idx = max((i for i,c in enumerate(children2) if c.tag == W + "r"), default=len(children2)-1)
    p_elem.insert(last_r_idx+1, end)
    # Add commentReference run after end
    r_ref = ET.Element(W + "r")
    # minimal rPr
    ref_inner = ET.Element(W + "commentReference"); ref_inner.set(W + "id", cid)
    r_ref.append(ref_inner)
    p_elem.append(r_ref)

stored["word/document.xml"] = ET.tostring(root, encoding="unicode").encode("utf-8")

# Create word/comments.xml
comments_root = ET.Element(W + "comments", {"xmlns:w": W_NS.strip("{}"), "xmlns:wp": WP, "xmlns:r": R_NS})
for cid in comment_ids:
    comm = ET.SubElement(comments_root, W + "comment", {W + "id": cid, W + "author":"Tester", W + "date":"2026-08-31T00:00:00Z", W + "initials":"T"})
    ET.SubElement(comm, W + "p").append(ET.fromstring(f'<w:r xmlns:w="{W_NS.strip("{}")}"><w:t>Comment {cid} body</w:t></w:r>'))
# Add one extra for control not needed

# Need proper namespace handling: create string manually
import xml.etree.ElementTree as ET2
# Instead write raw xml string for simplicity
comments_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="{W_NS.strip("{}")}" xmlns:r="{R_NS}" xmlns:wp="{WP}" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w15="http://schemas.microsoft.com/office/2011/word">
'''
for cid in comment_ids:
    comments_xml += f'<w:comment w:id="{cid}" w:author="Tester" w:date="2026-08-31T00:00:00Z" w:initials="T"><w:p><w:r><w:t>Comment {cid} body for {labels[int(cid)-10].lower()}</w:t></w:r></w:p></w:comment>\n'
comments_xml += '</w:comments>'

stored["word/comments.xml"] = comments_xml.encode("utf-8")

# Update [Content_Types].xml to include comments
ct = stored.get("[Content_Types].xml", b"").decode("utf-8")
if "comments.xml" not in ct:
    ct = ct.replace("</Types>", '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/></Types>')
    stored["[Content_Types].xml"] = ct.encode("utf-8")

# Fix relationships if needed: ensure word/_rels/document.xml.rels has comments entry (optional but nice)
rels_xml = stored.get("word/_rels/document.xml.rels", b"").decode("utf-8")
if "comments" not in rels_xml:
    # Inject Relationship
    rels_root = ET.fromstring(rels_xml if rels_xml else '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>')
    NS_OPC = "http://schemas.openxmlformats.org/package/2006/relationships"
    # find max Id
    ids = [r.get("Id") for r in rels_root.findall(f"{{{NS_OPC}}}Relationship")]
    new_id = "rIdComm1"
    counter=1
    while new_id in ids:
        counter+=1; new_id=f"rIdComm{counter}"
    ET.SubElement(rels_root, f"{{{NS_OPC}}}Relationship", {"Id": new_id, "Type":"http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments", "Target":"comments.xml"})
    stored["word/_rels/document.xml.rels"] = ET.tostring(rels_root, encoding="unicode").encode("utf-8")

with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z2:
    for fname, data in stored.items():
        z2.writestr(fname, data)

print("Fixture written with comment ranges", out)
# Verify quick parse
from adapter.ooxml_parser import OoxmlParser
p = OoxmlParser(out)
blocks = p.parse_document()
from core.model import Image, CommentRangeStart, CommentRangeEnd, NoteReference
cnt_start = sum(1 for b in blocks if hasattr(b,'content') for c in getattr(b,'content',[]) if isinstance(c, CommentRangeStart))
cnt_end = sum(1 for b in blocks if hasattr(b,'content') for c in getattr(b,'content',[]) if isinstance(c, CommentRangeEnd))
imgs = []
for b in blocks:
    if hasattr(b,'content'):
        for c in b.content:
            if isinstance(c, Image): imgs.append(c)
print(f"parsed starts {cnt_start} ends {cnt_end} images {len(imgs)}")
for im in imgs:
    print(im.image_id, im.wrap_type, im.wrap_mode, im.relative_from_horizontal, im.offset_horizontal, im.extent_cx, im.wrap_distances, im.wrap_polygon is not None)
comments = p.get_comments()
print(f"comments {len(comments)} ids {[c.note_id for c in comments]}")
