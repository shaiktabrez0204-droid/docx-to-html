"""Generate REAL .docx fixtures that actually carry w:numPr + numbering.xml.

These are genuine OOXML packages (not faked HTML): heading levels come from the
HeadingN style outline, and numbering comes from real w:numPr paragraph
properties referencing a real word/numbering.xml. The visible text never carries
the authoritative number.
"""

import os
import shutil
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIXTURES = os.path.join(PROJECT_ROOT, "tests", "fixtures")

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
RELS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT = "http://schemas.openxmlformats.org/package/2006/content-types"

NUMBERING_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="0">
    <w:multiLevelType w:val="hybridMultilevel"/>
    <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/><w:lvlJc w:val="left"/></w:lvl>
    <w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1.%2."/><w:lvlJc w:val="left"/></w:lvl>
    <w:lvl w:ilvl="2"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1.%2.%3."/><w:lvlJc w:val="left"/></w:lvl>
    <w:lvl w:ilvl="3"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1.%2.%3.%4."/><w:lvlJc w:val="left"/></w:lvl>
    <w:lvl w:ilvl="4"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1.%2.%3.%4.%5."/><w:lvlJc w:val="left"/></w:lvl>
    <w:lvl w:ilvl="5"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1.%2.%3.%4.%5.%6."/><w:lvlJc w:val="left"/></w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
</w:numbering>
"""


def _add_numpr(paragraph, num_id, ilvl):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def _inject_numbering(docx_path):
    """Add word/numbering.xml + its relationship + content-type override."""
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}

    # relationship id not already used
    rels = data.get("word/_rels/document.xml.rels", b"")
    rels_root = ET.fromstring(rels)
    used = [r.get("Id") for r in rels_root.findall("{%s}Relationship" % RELS)]
    n = 1
    while ("rIdN%d" % n) in used:
        n += 1
    rid = "rIdN%d" % n
    new_rel = ET.SubElement(rels_root, "{%s}Relationship" % RELS)
    new_rel.set("Id", rid)
    new_rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering")
    new_rel.set("Target", "numbering.xml")
    rels_out = ET.tostring(rels_root, encoding="UTF-8", xml_declaration=True)

    # content type override
    ct = data.get("[Content_Types].xml", b"")
    ct_root = ET.fromstring(ct)
    has = any(o.get("PartName") == "/word/numbering.xml"
              for o in ct_root.findall("{%s}Override" % CT))
    if not has:
        ov = ET.SubElement(ct_root, "{%s}Override" % CT)
        ov.set("PartName", "/word/numbering.xml")
        ov.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml")
    ct_out = ET.tostring(ct_root, encoding="UTF-8", xml_declaration=True)

    data["word/numbering.xml"] = NUMBERING_XML.encode("utf-8")
    data["word/_rels/document.xml.rels"] = rels_out
    data["[Content_Types].xml"] = ct_out

    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, b in data.items():
            zout.writestr(n, b)
    shutil.move(tmp, docx_path)


def _heading(doc, text, level, num=None):
    p = doc.add_heading(text, level=level)
    if num is not None:
        _add_numpr(p, num[0], num[1])
    return p


def _build(filename, builder):
    path = os.path.join(FIXTURES, filename)
    doc = Document()
    builder(doc)
    doc.save(path)
    _inject_numbering(path)
    print("WROTE", filename)


# 1. H1 numbered
def f1(doc):
    _heading(doc, "Introduction", 1, num=(1, 0))
    doc.add_paragraph("Some body text.")


# 2. H1 -> H2 numbered
def f2(doc):
    _heading(doc, "Chapter One", 1, num=(1, 0))
    _heading(doc, "Section A", 2, num=(1, 1))
    _heading(doc, "Section B", 2, num=(1, 1))


# 3. H1 -> H2 -> H3 numbered
def f3(doc):
    _heading(doc, "Top", 1, num=(1, 0))
    _heading(doc, "Mid", 2, num=(1, 1))
    _heading(doc, "Bottom", 3, num=(1, 2))


# 4. H1 -> H3 skipped level (numbered)
def f4(doc):
    _heading(doc, "Part One", 1, num=(1, 0))
    _heading(doc, "Deep Detail", 3, num=(1, 2))
    _heading(doc, "Section Two", 2, num=(1, 1))


# 5. numbering restart (Chapter pattern)
def f5(doc):
    _heading(doc, "Chapter One", 1, num=(1, 0))
    _heading(doc, "Intro", 2, num=(1, 1))
    _heading(doc, "Setup", 2, num=(1, 1))
    _heading(doc, "Chapter Two", 1, num=(1, 0))
    _heading(doc, "Design", 2, num=(1, 1))
    _heading(doc, "Build", 2, num=(1, 1))


# 6. duplicate heading text with numbering
def f6(doc):
    _heading(doc, "Introduction", 1, num=(1, 0))
    _heading(doc, "Details", 2, num=(1, 1))
    _heading(doc, "Introduction", 1, num=(1, 0))
    _heading(doc, "Details", 2, num=(1, 1))


# 7. unnumbered headings (numbering optional / absent)
def f7(doc):
    _heading(doc, "Overview", 1)
    _heading(doc, "Components", 2)
    _heading(doc, "Database", 3)


# 8. custom heading style + numbering
def f8(doc):
    styles = doc.styles
    base = styles["Heading1"]
    custom = styles.add_style("MyNumberedHeading", 1)  # 1 = paragraph style
    custom.base_style = base
    custom.name = "MyNumberedHeading"
    p = doc.add_paragraph("Custom Heading One", style=custom)
    _add_numpr(p, 1, 0)
    p2 = doc.add_paragraph("Custom Sub", style=styles["Heading2"])
    _add_numpr(p2, 1, 1)


# 9. inconsistent style/numbering combination
def f9(doc):
    # Heading2 style (level 2) but numbered as ilvl 0 (says H1) -> mismatch
    _heading(doc, "Looks Like Chapter", 2, num=(1, 0))
    _heading(doc, "Real Subsection", 3, num=(1, 2))


# 10. mixed numbered and unnumbered headings
def f10(doc):
    _heading(doc, "Numbered Chapter", 1, num=(1, 0))
    _heading(doc, "Numbered Section", 2, num=(1, 1))
    _heading(doc, "Unnumbered Part", 1)
    _heading(doc, "Unnumbered Sub", 2)


# 11. adversarial: "9. Fake Heading" text, NO numPr -> must not be authoritative
def f11(doc):
    _heading(doc, "9. Fake Heading", 2)  # Heading2 style, no numbering metadata
    _heading(doc, "Real Section", 3, num=(1, 2))


# 12. adversarial: numbering metadata present but text has no visible number
def f12(doc):
    _heading(doc, "Chapter Alpha", 1, num=(1, 0))  # text has no "1." but metadata does
    _heading(doc, "Section Beta", 2, num=(1, 1))


BUILDERS = {
    "num-h1.docx": f1,
    "num-h1-h2.docx": f2,
    "num-h1-h2-h3.docx": f3,
    "num-skipped.docx": f4,
    "num-restart.docx": f5,
    "num-duplicate.docx": f6,
    "num-unnumbered.docx": f7,
    "num-custom-style.docx": f8,
    "num-inconsistent.docx": f9,
    "num-mixed.docx": f10,
    "num-adversarial-text.docx": f11,
    "num-adversarial-metadata.docx": f12,
}


def main():
    for fn, b in BUILDERS.items():
        _build(fn, b)


if __name__ == "__main__":
    main()
