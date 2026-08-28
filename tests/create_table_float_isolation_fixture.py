import os, io, zipfile, xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm
from PIL import Image as PILImage

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIX = os.path.join(PROJ, "tests", "fixtures")
TMP = os.path.join(PROJ, "tests", ".table_flt_tmp")
os.makedirs(FIX, exist_ok=True)
os.makedirs(TMP, exist_ok=True)
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

def _pil_bytes(size, color):
    img = PILImage.new("RGB", size, color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()

def _img_file(name, data):
    p=os.path.join(TMP, name)
    with open(p,"wb") as f: f.write(data)
    return p

colors = [(200,30,30),(30,60,200),(30,160,60),(220,180,20)]
datas = [_pil_bytes((200,150),c) for c in colors]
files = [_img_file(f"tbl{i}.png", d) for i,d in enumerate(datas)]

def _convert_inline_to_anchor(docx_path, specs):
    with zipfile.ZipFile(docx_path,"r") as z:
        stored={i.filename: z.read(i.filename) for i in z.infolist()}
    root=ET.fromstring(stored["word/document.xml"])
    W="{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    WPq="{%s}"%WP
    inlines=list(root.iter(WPq+"inline"))
    assert len(inlines)==len(specs), f"expected {len(specs)} inlines got {len(inlines)}"
    for inline, spec in zip(inlines, specs):
        anchor=ET.Element(WPq+"anchor", {"distT":"0","distB":"0","distL":"0","distR":"0","simplePos":"0","relativeHeight":"251658240","behindDoc":"0","locked":"0","layoutInCell":"1","allowOverlap":"1"})
        for child in list(inline):
            anchor.append(child)
        sp=ET.SubElement(anchor, WPq+"simplePos"); sp.set("x","0"); sp.set("y","0")
        ph=ET.SubElement(anchor, WPq+"positionH"); ph.set("relativeFrom", spec["relH"])
        o=ET.SubElement(ph, WPq+"posOffset"); o.text=str(spec.get("offH",0))
        pv=ET.SubElement(anchor, WPq+"positionV"); pv.set("relativeFrom", spec["relV"])
        o2=ET.SubElement(pv, WPq+"posOffset"); o2.text=str(spec.get("offV",0))
        wrap=ET.SubElement(anchor, WPq+"wrapNone")
        parent=None
        for d in root.iter(W+"drawing"):
            if inline in list(d):
                parent=d; break
        idx=list(parent).index(inline)
        parent.remove(inline); parent.insert(idx, anchor)
    stored["word/document.xml"]=ET.tostring(root, encoding="unicode").encode("utf-8")
    with zipfile.ZipFile(docx_path,"w", zipfile.ZIP_DEFLATED) as z2:
        for fn, data in stored.items(): z2.writestr(fn, data)

doc=Document()
# H1 Intro
doc.add_heading("Introduction", level=1)
doc.add_paragraph("Intro paragraph.")
tbl=doc.add_table(rows=1, cols=1)
tbl.style="Table Grid"
cell=tbl.cell(0,0)
cell.text=""
p=cell.paragraphs[0]
r=p.add_run()
r.add_picture(files[0], width=Cm(4), height=Cm(3))
p.add_run(" Image A in table Intro")

# H2 Architecture
doc.add_heading("Architecture", level=2)
doc.add_paragraph("Architecture paragraph.")
tbl=doc.add_table(rows=1, cols=1)
tbl.style="Table Grid"
cell=tbl.cell(0,0)
cell.text=""
p=cell.paragraphs[0]
r=p.add_run()
r.add_picture(files[1], width=Cm(4), height=Cm(3))
p.add_run(" Image B in table Arch")

# H2 Implementation
doc.add_heading("Implementation", level=2)
doc.add_paragraph("Implementation paragraph.")
tbl=doc.add_table(rows=1, cols=1)
tbl.style="Table Grid"
cell=tbl.cell(0,0)
cell.text=""
p=cell.paragraphs[0]
r=p.add_run()
r.add_picture(files[2], width=Cm(4), height=Cm(3))
p.add_run(" Image C in table Impl")

# H1 Conclusion
doc.add_heading("Conclusion", level=1)
doc.add_paragraph("Conclusion paragraph.")
tbl=doc.add_table(rows=1, cols=1)
tbl.style="Table Grid"
cell=tbl.cell(0,0)
cell.text=""
p=cell.paragraphs[0]
r=p.add_run()
r.add_picture(files[3], width=Cm(4), height=Cm(3))
p.add_run(" Image D in table Concl")

out=os.path.join(FIX, "table-float-isolation.docx")
doc.save(out)
specs=[
    {"relH":"page","offH":800000,"relV":"page","offV":300000},
    {"relH":"page","offH":1500000,"relV":"page","offV":600000},
    {"relH":"page","offH":2200000,"relV":"page","offV":900000},
    {"relH":"page","offH":3800000,"relV":"page","offV":1200000},
]
_convert_inline_to_anchor(out, specs)
print("wrote", out)
