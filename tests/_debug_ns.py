import os, io, zipfile, xml.etree.ElementTree as ET
from docx import Document

TMP = "tests/.img_tmp"
d = Document()
p = d.add_paragraph()
p.add_run().add_picture(os.path.join(TMP, "r.png"))
buf = io.BytesIO()
d.save(buf)
z = zipfile.ZipFile(buf)
xml = z.read("word/document.xml").decode("utf-8")
root = ET.fromstring(xml)
wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
inline = root.find(".//" + "{" + wp + "}" + "inline")
print("find wp:inline (URI):", inline is not None)
print("root tag:", root.tag)
import re
m = re.search(r'xmlns:wp="([^"]+)"', xml)
print("declared wp ns:", m.group(1) if m else None)
# Try without namespace
print("contains '<wp:inline':", "<wp:inline" in xml)
