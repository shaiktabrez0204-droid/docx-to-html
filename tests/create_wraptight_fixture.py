import os, io, zipfile, xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm
from PIL import Image as PILImage

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIX = os.path.join(PROJ, "tests", "fixtures")
TMP = os.path.join(PROJ, "tests", ".wrap_tmp")
os.makedirs(FIX, exist_ok=True)
os.makedirs(TMP, exist_ok=True)
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

def pil_bytes(size, color):
    im = PILImage.new("RGB", size, color)
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return buf.getvalue()

def img_file(name, data):
    p = os.path.join(TMP, name)
    with open(p, "wb") as f:
        f.write(data)
    return p

RED = pil_bytes((180,130),(200,30,30))
BLUE = pil_bytes((180,130),(30,60,200))
GREEN = pil_bytes((180,130),(30,160,60))
YELL = pil_bytes((180,130),(220,180,20))
PURP = pil_bytes((180,130),(120,40,160))

files = {k: img_file(k+".png", v) for k,v in [("red",RED),("blue",BLUE),("green",GREEN),("yell",YELL),("purp",PURP)]}

def convert(docx_path, specs):
    with zipfile.ZipFile(docx_path, "r") as z:
        stored = {i.filename: z.read(i.filename) for i in z.infolist()}
    root = ET.fromstring(stored["word/document.xml"])
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    WPq = "{%s}" % WP
    inlines = list(root.iter(WPq + "inline"))
    assert len(inlines)==len(specs), f"expected {len(specs)} got {len(inlines)}"
    for inline, spec in zip(inlines, specs):
        anchor = ET.Element(WPq + "anchor", {
            "distT": str(spec.get("distT",0)),"distB": str(spec.get("distB",0)),"distL": str(spec.get("distL",0)),"distR": str(spec.get("distR",0)),
            "simplePos":"0","relativeHeight":"251658240","behindDoc":"0","locked":"0","layoutInCell":"1","allowOverlap":"1",
        })
        for child in list(inline):
            anchor.append(child)
        ET.SubElement(anchor, WPq + "simplePos", {"x":"0","y":"0"})
        ph = ET.SubElement(anchor, WPq + "positionH", {"relativeFrom": spec["relH"]})
        ET.SubElement(ph, WPq + "align").text = spec["alignH"]
        pv = ET.SubElement(anchor, WPq + "positionV", {"relativeFrom": spec["relV"]})
        ET.SubElement(pv, WPq + "align").text = spec.get("alignV","top")
        wrap_el = spec["wrap"]
        if wrap_el=="wrapTight":
            w = ET.SubElement(anchor, WPq + "wrapTight", {"wrapText":"bothSides","distT":str(spec.get("wrapDistT",0)),"distB":str(spec.get("wrapDistB",0))})
            # Add a simple rectangular polygon (square) for tight - 4 points
            poly = ET.SubElement(w, WPq + "wrapPolygon", {"edited":"0"})
            ET.SubElement(poly, WPq + "start", {"x":"0","y":"0"})
            ET.SubElement(poly, WPq + "lineTo", {"x":"21600","y":"0"})
            ET.SubElement(poly, WPq + "lineTo", {"x":"21600","y":"21600"})
            ET.SubElement(poly, WPq + "lineTo", {"x":"0","y":"21600"})
            ET.SubElement(poly, WPq + "lineTo", {"x":"0","y":"0"})
        elif wrap_el=="wrapThrough":
            w = ET.SubElement(anchor, WPq + "wrapThrough", {"wrapText":"bothSides","distT":str(spec.get("wrapDistT",0)),"distB":str(spec.get("wrapDistB",0))})
            poly = ET.SubElement(w, WPq + "wrapPolygon", {"edited":"0"})
            ET.SubElement(poly, WPq + "start", {"x":"0","y":"0"})
            ET.SubElement(poly, WPq + "lineTo", {"x":"21600","y":"0"})
            ET.SubElement(poly, WPq + "lineTo", {"x":"21600","y":"21600"})
            ET.SubElement(poly, WPq + "lineTo", {"x":"0","y":"21600"})
            ET.SubElement(poly, WPq + "lineTo", {"x":"0","y":"0"})
        elif wrap_el=="wrapTopAndBottom":
            w = ET.SubElement(anchor, WPq + "wrapTopAndBottom", {"distT":str(spec.get("wrapDistT",0)),"distB":str(spec.get("wrapDistB",0))})
            ET.SubElement(w, WPq + "effectExtent", {"l":"0","t":"0","r":"0","b":"0"})
        elif wrap_el=="wrapSquare":
            w = ET.SubElement(anchor, WPq + "wrapSquare", {"wrapText":"bothSides","distT":str(spec.get("wrapDistT",0)),"distB":str(spec.get("wrapDistB",0)),"distL":str(spec.get("wrapDistL",0)),"distR":str(spec.get("wrapDistR",0))})
            ET.SubElement(w, WPq + "effectExtent", {"l":"0","t":"0","r":"0","b":"0"})
        else:
            ET.SubElement(anchor, WPq + "wrapNone")
        parent=None
        for d in root.iter(W + "drawing"):
            if inline in list(d):
                parent=d; break
        idx=list(parent).index(inline)
        parent.remove(inline)
        parent.insert(idx, anchor)
    stored["word/document.xml"] = ET.tostring(root, encoding="unicode").encode("utf-8")
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as z2:
        for fn, data in stored.items():
            z2.writestr(fn, data)

doc=Document()
# H1 Intro + wrapTight left
doc.add_heading("Introduction", level=1)
doc.add_paragraph("Introduction paragraph with wrapTight left. " + "Lorem ipsum dolor sit amet consectetur adipiscing elit. "*4)
p=doc.add_paragraph()
p.add_run().add_picture(files["red"], width=Cm(4), height=Cm(3))
p.add_run(" wrapTight left dist 0.2in")

# H2 Arch + wrapThrough right
doc.add_heading("Architecture", level=2)
doc.add_paragraph("Architecture paragraph with wrapThrough right. " + "Lorem ipsum "*6)
p=doc.add_paragraph()
p.add_run().add_picture(files["blue"], width=Cm(4), height=Cm(3))
p.add_run(" wrapThrough right")

# H3 Data Model + topAndBottom center
doc.add_heading("Data Model", level=3)
doc.add_paragraph("Data Model paragraph with topAndBottom center. " + "Lorem ipsum "*6)
p=doc.add_paragraph()
p.add_run().add_picture(files["green"], width=Cm(4), height=Cm(3))
p.add_run(" topAndBottom center")

# H2 Impl + wrapTight right with dist
doc.add_heading("Implementation", level=2)
doc.add_paragraph("Implementation paragraph with wrapTight right dist. " + "Lorem ipsum "*6)
p=doc.add_paragraph()
p.add_run().add_picture(files["yell"], width=Cm(4), height=Cm(3))
p.add_run(" wrapTight right dist")

# H1 Concl + wrapThrough left
doc.add_heading("Conclusion", level=1)
doc.add_paragraph("Conclusion paragraph with wrapThrough left. " + "Lorem ipsum "*6)
p=doc.add_paragraph()
p.add_run().add_picture(files["purp"], width=Cm(4), height=Cm(3))
p.add_run(" wrapThrough left")

out=os.path.join(FIX, "wraptight-isolation.docx")
doc.save(out)
specs=[
    {"wrap":"wrapTight","relH":"page","alignH":"left","relV":"page","alignV":"top","distT":114300,"distB":114300,"distL":114300,"distR":114300,"wrapDistT":114300,"wrapDistB":114300},
    {"wrap":"wrapThrough","relH":"page","alignH":"right","relV":"page","alignV":"top","distT":114300,"distB":114300,"distL":114300,"distR":114300,"wrapDistT":114300,"wrapDistB":114300},
    {"wrap":"wrapTopAndBottom","relH":"page","alignH":"center","relV":"page","alignV":"top","distT":114300,"distB":114300,"wrapDistT":114300,"wrapDistB":114300},
    {"wrap":"wrapTight","relH":"page","alignH":"right","relV":"page","alignV":"top","distT":228600,"distB":228600,"distL":228600,"distR":228600,"wrapDistT":228600,"wrapDistB":228600},
    {"wrap":"wrapThrough","relH":"page","alignH":"left","relV":"page","alignV":"top","distT":114300,"distB":114300,"distL":114300,"distR":114300,"wrapDistT":114300,"wrapDistB":114300},
]
convert(out, specs)
print("wrote", out)
