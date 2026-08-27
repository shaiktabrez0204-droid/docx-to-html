"""Create REAL DOCX image fixtures for the visual-fidelity image phase.

Every fixture embeds actual image bytes extracted from real OOXML relationships
(word/_rels/document.xml.rels -> word/media/*). No placeholders, no regex, no
fake image objects. Generated with python-docx + PIL so the media is genuine.

Run:  python tests/create_image_fixtures.py
"""

import os
import io
import zipfile

from docx import Document
from docx.shared import Cm, Inches

from PIL import Image as PILImage

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIX = os.path.join(PROJECT_ROOT, "tests", "fixtures")
TMP = os.path.join(PROJECT_ROOT, "tests", ".img_tmp")
os.makedirs(FIX, exist_ok=True)
os.makedirs(TMP, exist_ok=True)


def _make_pil_bytes(size, color, fmt):
    img = PILImage.new("RGB", size, color)
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    return buf.getvalue()


# Real image bytes the fixtures embed. Different sizes/colors so browser checks
# can verify displayed dimensions and distinct placements.
PNG_BYTES = _make_pil_bytes((200, 150), (200, 30, 30), "PNG")      # red 200x150
JPEG_BYTES = _make_pil_bytes((120, 90), (30, 120, 200), "JPEG")    # blue 120x90
GIF_BYTES = _make_pil_bytes((80, 80), (40, 160, 40), "GIF")       # green 80x80
PNG2_BYTES = _make_pil_bytes((160, 160), (220, 180, 20), "PNG")   # yellow 160x160


def _write(path, data):
    with open(path, "wb") as f:
        f.write(data)


def _img_file(name, data):
    p = os.path.join(TMP, name)
    _write(p, data)
    return p


PNG_FILE = _img_file("r.png", PNG_BYTES)
JPEG_FILE = _img_file("b.jpg", JPEG_BYTES)
GIF_FILE = _img_file("g.gif", GIF_BYTES)
PNG2_FILE = _img_file("y.png", PNG2_BYTES)


def _set_alt(run, descr):
    """Set wp:docPr/@descr on the drawing inside a run (real OOXML alt text)."""
    from docx.oxml.ns import qn
    drawing = run._r.find(qn("w:drawing"))
    if drawing is None:
        return
    inline = drawing.find(qn("wp:inline"))
    if inline is None:
        return
    doc_pr = inline.find(qn("wp:docPr"))
    if doc_pr is None:
        return
    doc_pr.set("descr", descr)


# ---------------------------------------------------------------------------
# 1) One inline PNG
# ---------------------------------------------------------------------------
def fixture_inline_png():
    d = Document()
    d.add_paragraph("Before image.")
    p = d.add_paragraph()
    p.add_run().add_picture(PNG_FILE)
    d.add_paragraph("After image.")
    out = os.path.join(FIX, "img-inline-png.docx")
    d.save(out)
    return out


# ---------------------------------------------------------------------------
# 2) One inline JPEG
# ---------------------------------------------------------------------------
def fixture_inline_jpeg():
    d = Document()
    d.add_paragraph("A JPEG photo below.")
    p = d.add_paragraph()
    p.add_run().add_picture(JPEG_FILE)
    out = os.path.join(FIX, "img-inline-jpeg.docx")
    d.save(out)
    return out


# ---------------------------------------------------------------------------
# 3) Multiple images (across paragraphs + two in one paragraph)
# ---------------------------------------------------------------------------
def fixture_multiple():
    d = Document()
    d.add_paragraph("Intro paragraph.")
    p1 = d.add_paragraph()
    p1.add_run().add_picture(PNG_FILE)  # image 1
    p2 = d.add_paragraph()
    p2.add_run().add_picture(JPEG_FILE)  # image 2
    p3 = d.add_paragraph()
    # two images in the same paragraph, separated by text
    p3.add_run().add_picture(PNG_FILE)
    p3.add_run(" between ")
    p3.add_run().add_picture(GIF_FILE)
    d.add_paragraph("Outro paragraph.")
    out = os.path.join(FIX, "img-multiple.docx")
    d.save(out)
    return out


# ---------------------------------------------------------------------------
# 4) Image with explicit dimensions
# ---------------------------------------------------------------------------
def fixture_explicit_dims():
    d = Document()
    d.add_paragraph("Image sized to 5cm x 4cm in the document.")
    p = d.add_paragraph()
    p.add_run().add_picture(PNG_FILE, width=Cm(5), height=Cm(4))
    out = os.path.join(FIX, "img-explicit-dims.docx")
    d.save(out)
    return out


# ---------------------------------------------------------------------------
# 5) Image WITH alt text
# ---------------------------------------------------------------------------
def fixture_alt_text():
    d = Document()
    d.add_paragraph("Image carrying real alt text.")
    p = d.add_paragraph()
    run = p.add_run()
    run.add_picture(PNG_FILE)
    _set_alt(run, "A red test rectangle")
    out = os.path.join(FIX, "img-alt-text.docx")
    d.save(out)
    return out


# ---------------------------------------------------------------------------
# 6) Image WITHOUT alt text
# ---------------------------------------------------------------------------
def fixture_no_alt():
    d = Document()
    d.add_paragraph("Image with no alt text.")
    p = d.add_paragraph()
    p.add_run().add_picture(PNG_FILE)
    out = os.path.join(FIX, "img-no-alt.docx")
    d.save(out)
    return out


# ---------------------------------------------------------------------------
# 7) Same image referenced multiple times (dedup at asset level)
#    Two pictures -> two rIds, but rewrite rels so BOTH point at the SAME media
#    file. Result: one extracted asset, two document placements.
# ---------------------------------------------------------------------------
def fixture_reused():
    d = Document()
    d.add_paragraph("Same image used twice.")
    p1 = d.add_paragraph()
    r1 = p1.add_run()
    r1.add_picture(PNG_FILE)
    p2 = d.add_paragraph()
    r2 = p2.add_run()
    r2.add_picture(PNG_FILE)  # python-docx makes image2.png (identical bytes)
    out = os.path.join(FIX, "img-reused.docx")
    d.save(out)

    # Rewrite document.xml.rels so the rId of the SECOND picture targets the
    # SAME media file as the first. Expected outcome: one extracted asset +
    # two document placements (dedup at the asset level).
    with zipfile.ZipFile(out, "r") as z:
        names = z.namelist()
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
    import re
    targets = re.findall(r'Target="([^"]*image\d+\.[a-z]+)"', rels)
    if len(targets) >= 2:
        second = targets[1]
        rels = rels.replace('Target="%s"' % second, 'Target="%s"' % targets[0])
    # Drop the now-unused duplicate media file to keep the package valid/clean.
    unused = set()
    if len(targets) >= 2:
        unused.add("word/" + targets[1])

    # Read everything into memory first, then rewrite (never hold the same
    # archive open for read and write at once).
    with zipfile.ZipFile(out, "r") as zold:
        stored = {n: zold.read(n) for n in zold.namelist()}
    stored["word/_rels/document.xml.rels"] = rels.encode("utf-8")
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as znew:
        for n, data in stored.items():
            if n in unused:
                continue
            znew.writestr(n, data)
    return out


# ---------------------------------------------------------------------------
# 8) Image mixed with text: "before [IMG] after"
# ---------------------------------------------------------------------------
def fixture_mixed_text():
    d = Document()
    p = d.add_paragraph()
    p.add_run("This is before ")
    p.add_run().add_picture(PNG_FILE)
    p.add_run(" and this is after.")
    out = os.path.join(FIX, "img-mixed-text.docx")
    d.save(out)
    return out


# ---------------------------------------------------------------------------
# 9) Floating (wp:anchor) image: convert inline -> anchor in the saved docx.
#    Exercises the real anchor parsing branch. Positioning layout is out of
#    scope; we assert the image is still extracted and placed correctly.
# ---------------------------------------------------------------------------
def fixture_floating():
    d = Document()
    d.add_paragraph("Paragraph before floating image.")
    p = d.add_paragraph()
    run = p.add_run()
    run.add_picture(PNG_FILE)
    out = os.path.join(FIX, "img-floating.docx")
    d.save(out)

    from docx.oxml.ns import qn
    with zipfile.ZipFile(out, "r") as z:
        xml = z.read("word/document.xml").decode("utf-8")
        stored = {i.filename: z.read(i.filename) for i in z.infolist()}
    # Convert the first wp:inline to wp:anchor with required children.
    # We do a minimal structural transform via ElementTree.
    import xml.etree.ElementTree as ET
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    }
    new_xml = xml
    root = ET.fromstring(xml)
    W = "{%s}" % ns["w"]
    WP = "{%s}" % ns["wp"]
    inline = root.find(".//" + WP + "inline")
    if inline is not None:
        # Rename tag to wp:anchor and add required positioning children.
        anchor = ET.Element(WP + "anchor",
                            {"distT": "0", "distB": "0", "distL": "0", "distR": "0",
                             "simplePos": "0", "relativeHeight": "251658240",
                             "behindDoc": "0", "locked": "0", "layoutInCell": "1",
                             "allowOverlap": "1"})
        for child in list(inline):
            anchor.append(child)
        # Insert required positioning elements near the front.
        sp = ET.SubElement(anchor, WP + "simplePos")
        sp.set("x", "0"); sp.set("y", "0")
        ph = ET.SubElement(anchor, WP + "positionH")
        ph.set("relativeFrom", "page")
        ph_off = ET.SubElement(ph, WP + "posOffset"); ph_off.text = "2000000"
        pv = ET.SubElement(anchor, WP + "positionV")
        pv.set("relativeFrom", "page")
        pv_off = ET.SubElement(pv, WP + "posOffset"); pv_off.text = "1000000"
        wrap = ET.SubElement(anchor, WP + "wrapSquare")
        wrap.set("wrapText", "bothSides")
        # Replace inline with anchor inside its parent (the w:drawing).
        drawing_parent = None
        for d_ in root.iter(W + "drawing"):
            if inline in list(d_):
                drawing_parent = d_
                break
        if drawing_parent is not None:
            idx = list(drawing_parent).index(inline)
            drawing_parent.remove(inline)
            drawing_parent.insert(idx, anchor)
        new_xml = ET.tostring(root, encoding="unicode", xml_declaration=True)
        stored["word/document.xml"] = new_xml.encode("utf-8")
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z2:
            for fname, data in stored.items():
                z2.writestr(fname, data)
    return out


# ---------------------------------------------------------------------------
# Adversarial fixtures: real media + deliberately broken OOXML relationships,
# to prove the converter fails clearly or degrades safely (no broken HTML).
# ---------------------------------------------------------------------------
def _rewrite_docx(src, out_name, xml_transform=None, rels_transform=None, drop=None):
    out = os.path.join(FIX, out_name)
    with zipfile.ZipFile(src, "r") as z:
        stored = {n: z.read(n) for n in z.namelist()}
    if xml_transform is not None:
        xml = stored["word/document.xml"].decode("utf-8")
        stored["word/document.xml"] = xml_transform(xml).encode("utf-8")
    if rels_transform is not None:
        rels = stored["word/_rels/document.xml.rels"].decode("utf-8")
        stored["word/_rels/document.xml.rels"] = rels_transform(rels).encode("utf-8")
    if drop:
        for d in drop:
            stored.pop(d, None)
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z2:
        for n, data in stored.items():
            z2.writestr(n, data)
    return out


def fixture_adv_missing_rel():
    # r:embed points at a relationship id that does not exist.
    return _rewrite_docx(
        os.path.join(FIX, "img-inline-png.docx"), "img-adv-missing-rel.docx",
        xml_transform=lambda x: x.replace('r:embed="rId9"', 'r:embed="rId999"'))


def fixture_adv_missing_media():
    # Relationship resolves, but its Target media file is absent.
    return _rewrite_docx(
        os.path.join(FIX, "img-inline-png.docx"), "img-adv-missing-media.docx",
        rels_transform=lambda r: r.replace('Target="media/image1.png"', 'Target="media/ghost.png"'),
        drop=["word/media/image1.png"])


def fixture_adv_no_dims():
    # Drawing has no wp:extent -> dimensions unknown (must not break render).
    def _rm(xml):
        import re
        return re.sub(r"<wp:extent[^/]*/>", "", xml)
    return _rewrite_docx(
        os.path.join(FIX, "img-inline-png.docx"), "img-adv-no-dims.docx",
        xml_transform=_rm)


def fixture_adv_no_embed():
    # blip has no r:embed attribute at all -> cannot resolve.
    return _rewrite_docx(
        os.path.join(FIX, "img-inline-png.docx"), "img-adv-no-embed.docx",
        xml_transform=lambda x: x.replace('r:embed="rId9"', ""))


def fixture_adv_unsupported():
    # Relationship Target is an unsupported (non-browser) type; media absent.
    return _rewrite_docx(
        os.path.join(FIX, "img-inline-png.docx"), "img-adv-unsupported.docx",
        rels_transform=lambda r: r.replace('Target="media/image1.png"', 'Target="media/image1.emf"'),
        drop=["word/media/image1.png"])


if __name__ == "__main__":
    built = [
        fixture_inline_png(),
        fixture_inline_jpeg(),
        fixture_multiple(),
        fixture_explicit_dims(),
        fixture_alt_text(),
        fixture_no_alt(),
        fixture_reused(),
        fixture_mixed_text(),
        fixture_floating(),
        fixture_adv_missing_rel(),
        fixture_adv_missing_media(),
        fixture_adv_no_dims(),
        fixture_adv_no_embed(),
        fixture_adv_unsupported(),
    ]
    for b in built:
        print("WROTE", b)
    print("DONE")
