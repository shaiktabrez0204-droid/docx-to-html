"""Create REAL DOCX floating-image fixtures for the anchored-layout phase.

Every fixture embeds genuine image bytes (python-docx + PIL) and a real wp:anchor
element with explicit OOXML positioning metadata (positionH/positionV/wrap/dist).
No regex, no fake image objects. The anchor is produced by transforming the
inline drawing that python-docx emits, exactly like create_image_fixtures.py.

Run:  python tests/create_anchor_fixtures.py
"""

import os
import io
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.shared import Cm
from PIL import Image as PILImage

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIX = os.path.join(PROJECT_ROOT, "tests", "fixtures")
TMP = os.path.join(PROJECT_ROOT, "tests", ".flt_tmp")
os.makedirs(FIX, exist_ok=True)
os.makedirs(TMP, exist_ok=True)

WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def _pil_bytes(size, color, fmt="PNG"):
    img = PILImage.new("RGB", size, color)
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    return buf.getvalue()


# Real image bytes at varied sizes so display geometry is verifiable.
RED = _pil_bytes((300, 200), (200, 30, 30))
BLUE = _pil_bytes((160, 120), (30, 60, 200))
GREEN = _pil_bytes((100, 100), (30, 160, 60))
YELLOW = _pil_bytes((260, 180), (220, 180, 20))
SMALL = _pil_bytes((80, 60), (120, 40, 160))
LARGE = _pil_bytes((420, 320), (20, 120, 120))


def _img_file(name, data):
    p = os.path.join(TMP, name)
    with open(p, "wb") as f:
        f.write(data)
    return p


RED_F = _img_file("fr.png", RED)
BLUE_F = _img_file("fb.png", BLUE)
GREEN_F = _img_file("fg.png", GREEN)
YELLOW_F = _img_file("fy.png", YELLOW)
SMALL_F = _img_file("fs.png", SMALL)
LARGE_F = _img_file("fl.png", LARGE)


def _convert_inline_to_anchor(docx_path, specs):
    """Open docx_path, convert its inline drawings (in order) to wp:anchors.

    ``specs`` is a list (one per inline drawing) of dicts:
        relH, relV, alignH, alignV, offH (EMU), offV (EMU),
        wrap ('Square'|'None'|'TopAndBottom'), distT/B/L/R (EMU),
        drop_extent (bool), behind (bool)
    Missing positionH/positionV when the corresponding dict key is absent.
    """
    with zipfile.ZipFile(docx_path, "r") as z:
        stored = {i.filename: z.read(i.filename) for i in z.infolist()}
    root = ET.fromstring(stored["word/document.xml"])
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    WPq = "{%s}" % WP

    inlines = list(root.iter(WPq + "inline"))
    assert len(inlines) == len(specs), \
        "expected %d inline drawings, found %d" % (len(specs), len(inlines))

    for inline, spec in zip(inlines, specs):
        anchor = ET.Element(WPq + "anchor", {
            "distT": str(spec.get("distT", 0)),
            "distB": str(spec.get("distB", 0)),
            "distL": str(spec.get("distL", 0)),
            "distR": str(spec.get("distR", 0)),
            "simplePos": "0",
            "relativeHeight": "251658240",
            "behindDoc": "1" if spec.get("behind") else "0",
            "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
        })
        for child in list(inline):
            anchor.append(child)
        if spec.get("drop_extent"):
            for c in list(anchor):
                if c.tag == WPq + "extent":
                    anchor.remove(c)
        sp = ET.SubElement(anchor, WPq + "simplePos")
        sp.set("x", "0"); sp.set("y", "0")

        if "relH" in spec:
            ph = ET.SubElement(anchor, WPq + "positionH")
            ph.set("relativeFrom", spec["relH"])
            if spec.get("alignH"):
                a = ET.SubElement(ph, WPq + "align"); a.text = spec["alignH"]
            else:
                o = ET.SubElement(ph, WPq + "posOffset")
                o.text = str(spec.get("offH", 0))
        if "relV" in spec:
            pv = ET.SubElement(anchor, WPq + "positionV")
            pv.set("relativeFrom", spec["relV"])
            if spec.get("alignV"):
                a = ET.SubElement(pv, WPq + "align"); a.text = spec["alignV"]
            else:
                o = ET.SubElement(pv, WPq + "posOffset")
                o.text = str(spec.get("offV", 0))

        wk = spec.get("wrap", "Square")
        wrap = ET.SubElement(anchor, WPq + "wrap" + wk)
        if wk == "Square":
            wrap.set("wrapText", "bothSides")

        # Replace inline with anchor inside its w:drawing parent.
        parent = None
        for d in root.iter(W + "drawing"):
            if inline in list(d):
                parent = d
                break
        idx = list(parent).index(inline)
        parent.remove(inline)
        parent.insert(idx, anchor)

    stored["word/document.xml"] = ET.tostring(root, encoding="unicode").encode("utf-8")
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as z2:
        for fname, data in stored.items():
            z2.writestr(fname, data)


def _save(d, name):
    out = os.path.join(FIX, name)
    d.save(out)
    return out


# ---------------------------------------------------------------------------
# 1) page-relative, centered (align center both axes), wrapSquare
# ---------------------------------------------------------------------------
def fixture_page_center():
    d = Document()
    d.add_paragraph("Document title paragraph above the centered floating image.")
    p = d.add_paragraph()
    p.add_run().add_picture(YELLOW_F, width=Cm(6), height=Cm(4))
    out = _save(d, "flt-page-center.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "page", "alignH": "center",
        "relV": "page", "alignV": "center",
        "wrap": "Square", "distT": 0, "distB": 0, "distL": 0, "distR": 0,
    }])
    return out


# ---------------------------------------------------------------------------
# 2) page-relative, explicit offset (posOffset), wrapNone
# ---------------------------------------------------------------------------
def fixture_page_offset():
    d = Document()
    d.add_paragraph("Paragraph before the page-offset floating image.")
    p = d.add_paragraph()
    p.add_run().add_picture(RED_F, width=Cm(5), height=Cm(3.5))
    out = _save(d, "flt-page-offset.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "page", "offH": 2000000,   # ~210 px from page left
        "relV": "page", "offV": 1200000,   # ~126 px from page top
        "wrap": "None", "distT": 0, "distB": 0, "distL": 0, "distR": 0,
    }])
    return out


# ---------------------------------------------------------------------------
# 3) margin-relative, offset, wrapSquare
# ---------------------------------------------------------------------------
def fixture_margin():
    d = Document()
    d.add_paragraph("Margin-relative floating image below.")
    p = d.add_paragraph()
    p.add_run().add_picture(BLUE_F, width=Cm(4), height=Cm(3))
    out = _save(d, "flt-margin.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "margin", "offH": 1000000,  # ~105 px into the margin box
        "relV": "margin", "offV": 800000,   # ~84 px into the margin box
        "wrap": "Square", "distT": 50000, "distB": 50000, "distL": 50000, "distR": 50000,
    }])
    return out


# ---------------------------------------------------------------------------
# 4) paragraph-relative, offset, wrapSquare
# ---------------------------------------------------------------------------
def fixture_paragraph():
    d = Document()
    d.add_paragraph("A paragraph that precedes the anchored image.")
    p = d.add_paragraph()
    p.add_run("Anchored text before the image. ")
    p.add_run().add_picture(GREEN_F, width=Cm(3), height=Cm(3))
    p.add_run(" Anchored text after the image.")
    out = _save(d, "flt-paragraph.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "paragraph", "offH": 500000,  # ~53 px from paragraph left
        "relV": "paragraph", "offV": 0,
        "wrap": "Square", "distT": 0, "distB": 0, "distL": 0, "distR": 0,
    }])
    return out


# ---------------------------------------------------------------------------
# 5) column-relative, offset, wrapSquare (single column => margin box)
# ---------------------------------------------------------------------------
def fixture_column():
    d = Document()
    d.add_paragraph("Column-relative floating image.")
    p = d.add_paragraph()
    p.add_run().add_picture(BLUE_F, width=Cm(4), height=Cm(3))
    out = _save(d, "flt-column.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "column", "offH": 600000,
        "relV": "column", "offV": 400000,
        "wrap": "Square", "distT": 0, "distB": 0, "distL": 0, "distR": 0,
    }])
    return out


# ---------------------------------------------------------------------------
# 6) wrapSquare with margin-relative + left align => REAL CSS float (text wraps)
# ---------------------------------------------------------------------------
def fixture_wrapsquare():
    d = Document()
    d.add_paragraph("A paragraph with enough text so that a floated image on the "
                    "left has body copy flowing around it on the right side. " * 6)
    p = d.add_paragraph()
    p.add_run().add_picture(RED_F, width=Cm(5), height=Cm(3.5))
    out = _save(d, "flt-wrapsquare.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "margin", "alignH": "left",
        "relV": "margin", "offV": 0,
        "wrap": "Square", "distT": 50000, "distB": 50000, "distL": 50000, "distR": 50000,
    }])
    return out


# ---------------------------------------------------------------------------
# 7) wrapTopAndBottom (page-relative offset; absolute, documented approximation)
# ---------------------------------------------------------------------------
def fixture_wraptopbottom():
    d = Document()
    d.add_paragraph("Text before the top-and-bottom wrapped image.")
    p = d.add_paragraph()
    p.add_run().add_picture(YELLOW_F, width=Cm(6), height=Cm(4))
    d.add_paragraph("More body text after the image to show flow around it.")
    out = _save(d, "flt-wraptopbottom.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "page", "offH": 1500000,
        "relV": "page", "offV": 1000000,
        "wrap": "TopAndBottom", "distT": 0, "distB": 0, "distL": 0, "distR": 0,
    }])
    return out


# ---------------------------------------------------------------------------
# 8) wrapNone (page-relative offset; overlay, overlap acceptable by design)
# ---------------------------------------------------------------------------
def fixture_wrapnone():
    d = Document()
    d.add_paragraph("Body text that the wrapNone image may sit over.")
    p = d.add_paragraph()
    p.add_run().add_picture(RED_F, width=Cm(5), height=Cm(3.5))
    out = _save(d, "flt-wrapnone.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "page", "offH": 1200000,
        "relV": "page", "offV": 600000,
        "wrap": "None", "distT": 0, "distB": 0, "distL": 0, "distR": 0,
    }])
    return out


# ---------------------------------------------------------------------------
# 9) multiple floating images in one document
# ---------------------------------------------------------------------------
def fixture_multi():
    d = Document()
    d.add_paragraph("First floating image paragraph.")
    p1 = d.add_paragraph()
    p1.add_run().add_picture(RED_F, width=Cm(4), height=Cm(3))
    d.add_paragraph("Second floating image paragraph.")
    p2 = d.add_paragraph()
    p2.add_run().add_picture(BLUE_F, width=Cm(4), height=Cm(3))
    d.add_paragraph("Third floating image paragraph.")
    p3 = d.add_paragraph()
    p3.add_run().add_picture(GREEN_F, width=Cm(3), height=Cm(3))
    out = _save(d, "flt-multi.docx")
    _convert_inline_to_anchor(out, [
        {"relH": "page", "offH": 1000000, "relV": "page", "offV": 300000, "wrap": "None"},
        {"relH": "page", "offH": 3000000, "relV": "page", "offV": 300000, "wrap": "None"},
        {"relH": "page", "offH": 5500000, "relV": "page", "offV": 300000, "wrap": "None"},
    ])
    return out


# ---------------------------------------------------------------------------
# 10) floating image near a heading
# ---------------------------------------------------------------------------
def fixture_near_heading():
    d = Document()
    d.add_heading("Section Heading", level=1)
    p = d.add_paragraph()
    p.add_run().add_picture(YELLOW_F, width=Cm(5), height=Cm(3.5))
    d.add_paragraph("Body text following the heading and the floating image.")
    out = _save(d, "flt-near-heading.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "page", "offH": 2000000, "relV": "page", "offV": 1500000,
        "wrap": "Square",
    }])
    return out


# ---------------------------------------------------------------------------
# 11) floating image between two paragraphs (nearest-block association)
# ---------------------------------------------------------------------------
def fixture_near_paragraph():
    d = Document()
    d.add_paragraph("Paragraph A before the anchored image.")
    p = d.add_paragraph()
    p.add_run().add_picture(GREEN_F, width=Cm(3), height=Cm(3))
    d.add_paragraph("Paragraph B after the anchored image.")
    out = _save(d, "flt-near-paragraph.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "page", "offH": 2500000, "relV": "page", "offV": 800000,
        "wrap": "Square",
    }])
    return out


# ---------------------------------------------------------------------------
# 12) multiple sections (section break), one float each
# ---------------------------------------------------------------------------
def fixture_multi_section():
    d = Document()
    d.add_paragraph("Section one text.")
    p1 = d.add_paragraph()
    p1.add_run().add_picture(RED_F, width=Cm(4), height=Cm(3))
    d.add_section()  # section break
    d.add_paragraph("Section two text.")
    p2 = d.add_paragraph()
    p2.add_run().add_picture(BLUE_F, width=Cm(4), height=Cm(3))
    out = _save(d, "flt-multi-section.docx")
    _convert_inline_to_anchor(out, [
        {"relH": "page", "offH": 800000, "relV": "page", "offV": 300000, "wrap": "None"},
        {"relH": "page", "offH": 3500000, "relV": "page", "offV": 300000, "wrap": "None"},
    ])
    return out


# ---------------------------------------------------------------------------
# 13) different image sizes as floats
# ---------------------------------------------------------------------------
def fixture_sizes():
    d = Document()
    d.add_paragraph("Small floating image.")
    ps = d.add_paragraph()
    ps.add_run().add_picture(SMALL_F, width=Cm(2), height=Cm(1.5))
    d.add_paragraph("Large floating image.")
    pl = d.add_paragraph()
    pl.add_run().add_picture(LARGE_F, width=Cm(8), height=Cm(6))
    out = _save(d, "flt-sizes.docx")
    _convert_inline_to_anchor(out, [
        {"relH": "page", "offH": 500000, "relV": "page", "offV": 200000, "wrap": "None"},
        {"relH": "page", "offH": 3000000, "relV": "page", "offV": 3500000, "wrap": "None"},
    ])
    return out


# ---------------------------------------------------------------------------
# 14) low-confidence / ambiguous: anchor in an EMPTY paragraph between text
# ---------------------------------------------------------------------------
def fixture_lowconf():
    d = Document()
    d.add_paragraph("Paragraph A with real text content.")
    pe = d.add_paragraph()  # EMPTY paragraph holding only the drawing
    pe.add_run().add_picture(GREEN_F, width=Cm(3), height=Cm(3))
    d.add_paragraph("Paragraph B with real text content.")
    out = _save(d, "flt-lowconf.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "page", "offH": 2500000, "relV": "page", "offV": 800000,
        "wrap": "Square",
    }])
    return out


# ---------------------------------------------------------------------------
# Adversarial fixtures
# ---------------------------------------------------------------------------
def fixture_adv_no_posh():
    d = Document()
    d.add_paragraph("Anchor missing positionH.")
    p = d.add_paragraph()
    p.add_run().add_picture(RED_F, width=Cm(4), height=Cm(3))
    out = _save(d, "flt-adv-no-posh.docx")
    _convert_inline_to_anchor(out, [{
        "relV": "page", "offV": 500000, "wrap": "Square",
    }])
    return out


def fixture_adv_no_extent():
    d = Document()
    d.add_paragraph("Anchor with no wp:extent.")
    p = d.add_paragraph()
    p.add_run().add_picture(RED_F, width=Cm(4), height=Cm(3))
    out = _save(d, "flt-adv-no-extent.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "page", "offH": 500000, "relV": "page", "offV": 500000,
        "wrap": "Square", "drop_extent": True,
    }])
    return out


def fixture_adv_unsupported_wrap():
    d = Document()
    d.add_paragraph("Anchor with unsupported wrapTight.")
    p = d.add_paragraph()
    p.add_run().add_picture(RED_F, width=Cm(4), height=Cm(3))
    out = _save(d, "flt-adv-unsupported-wrap.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "page", "offH": 500000, "relV": "page", "offV": 500000,
        "wrap": "Tight",
    }])
    return out


def fixture_adv_unsupported_relfrom():
    d = Document()
    d.add_paragraph("Anchor with a relativeFrom value this engine does not model.")
    p = d.add_paragraph()
    p.add_run().add_picture(RED_F, width=Cm(4), height=Cm(3))
    out = _save(d, "flt-adv-unsupported-relfrom.docx")
    _convert_inline_to_anchor(out, [{
        "relH": "bogusCoord", "offH": 50000, "relV": "page", "offV": 500000,
        "wrap": "Square",
    }])
    return out


def fixture_adv_zero_dims():
    d = Document()
    d.add_paragraph("Anchor with zero extent.")
    p = d.add_paragraph()
    p.add_run().add_picture(RED_F, width=Cm(4), height=Cm(3))
    out = _save(d, "flt-adv-zero-dims.docx")
    # Convert then force extent cx/cy to 0.
    _convert_inline_to_anchor(out, [{
        "relH": "page", "offH": 500000, "relV": "page", "offV": 500000,
        "wrap": "Square",
    }])
    # Patch extent to 0/0.
    with zipfile.ZipFile(out, "r") as z:
        stored = {i.filename: z.read(i.filename) for i in z.infolist()}
    root = ET.fromstring(stored["word/document.xml"])
    WPq = "{%s}" % WP
    for ext in root.iter(WPq + "extent"):
        ext.set("cx", "0"); ext.set("cy", "0")
    stored["word/document.xml"] = ET.tostring(root, encoding="unicode").encode("utf-8")
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z2:
        for fn, data in stored.items():
            z2.writestr(fn, data)
    return out


if __name__ == "__main__":
    built = [
        fixture_page_center(), fixture_page_offset(), fixture_margin(),
        fixture_paragraph(), fixture_column(), fixture_wrapsquare(),
        fixture_wraptopbottom(), fixture_wrapnone(), fixture_multi(),
        fixture_near_heading(), fixture_near_paragraph(), fixture_multi_section(),
        fixture_sizes(), fixture_lowconf(),
        fixture_adv_no_posh(), fixture_adv_no_extent(),
        fixture_adv_unsupported_wrap(), fixture_adv_unsupported_relfrom(),
        fixture_adv_zero_dims(),
    ]
    for b in built:
        print("WROTE", b)
    print("DONE")
