import os, sys, zipfile, xml.etree.ElementTree as ET
sys.path.insert(0, '.')
from docx import Document
from docx.shared import Inches
from PIL import Image as PILImage
os.makedirs('C:/Temp/gen2', exist_ok=True)
im=PILImage.new('RGB',(200,100),(180,50,50)); im.save('C:/Temp/gen2/tb.png')
im2=PILImage.new('RGB',(150,80),(50,180,50)); im2.save('C:/Temp/gen2/tb2.png')

def create_doc(path):
    doc = Document()
    # Section 0: letter 1in margins, header distance 0.5in
    s0 = doc.sections[0]
    s0.page_width=Inches(8.5); s0.page_height=Inches(11)
    s0.left_margin=Inches(1); s0.right_margin=Inches(1); s0.top_margin=Inches(1); s0.bottom_margin=Inches(1)
    s0.header_distance=Inches(0.5); s0.footer_distance=Inches(0.5)
    # Body control topAndBottom
    # Add body paragraph with topAndBottom will be patched later via zip
    doc.add_paragraph('Body start')
    p = doc.add_paragraph('Body paragraph before')
    # Add section break for sec1 with different page size
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p2 = doc.add_paragraph('Section1 start')
    pPr = p2._element.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz'); pgSz.set(qn('w:w'),'11906'); pgSz.set(qn('w:h'),'16838'); sectPr.append(pgSz) # A4
    pgMar = OxmlElement('w:pgMar'); pgMar.set(qn('w:left'),'1440'); pgMar.set(qn('w:right'),'720'); pgMar.set(qn('w:top'),'720'); pgMar.set(qn('w:bottom'),'720'); pgMar.set(qn('w:header'),'360'); pgMar.set(qn('w:footer'),'360'); sectPr.append(pgMar)
    pPr.append(sectPr)
    doc.add_paragraph('Body section1 para')
    # Save base
    doc.save(path)
    return path

base='C:/Temp/gen2/base_tb.docx'
create_doc(base)
print('base saved', base)

# Now patch headers/footers with topAndBottom images via zip manipulation
import zipfile, xml.etree.ElementTree as ET

def patch_hf_with_tb(docx_path, out_path):
    z=zipfile.ZipFile(docx_path,'r')
    parts={n:z.read(n) for n in z.namelist()}
    z.close()
    # Ensure header1.xml, footer1.xml exist - create if not
    # Read document.xml to find header refs
    doc_xml = parts['word/document.xml'].decode()
    print(doc_xml[:2000])
    # For simplicity, ensure header1.xml and footer1.xml exist with minimal content then replace
    # We'll create new header1.xml with 3 topAndBottom images covering cases
    NS_W='http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    # Create header1.xml content
    # Need to handle 2 sections: section0 has no header ref, section1 final has header ref rId? In our base, final sectPr has no headerReference yet, only earlier? Let's check.
    # Actually we didn't add header refs. We'll need to add header references to sectPrs and create header parts.
    # Simpler: use low-level to add header parts and relationships via python-docx API then patch anchor XML.
    pass

# Use alternative: directly use python-docx to add header/footer then patch anchors
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def add_tb_to_header(doc_path, out_path):
    doc = Document(doc_path)
    # Ensure headers for both sections
    # Section 0 header
    s0 = doc.sections[0]
    # Section 1 if exists
    secs = doc.sections
    print('num sections', len(secs))
    for idx, sec in enumerate(secs):
        print(f'sec {idx} header distance {sec.header_distance} footer {sec.footer_distance} left {sec.left_margin}')
        hdr = sec.header
        hdr.is_linked_to_previous=False
        # clear
        for p in hdr.paragraphs:
            p.text=''
        # Add paragraph for header
        p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        p.text = f'Header sec{idx} before '
        # We will patch this paragraph to add topAndBottom anchors via xml after save
    s1 = secs[1] if len(secs)>1 else None
    if s1:
        hdr1 = s1.header
        hdr1.is_linked_to_previous=False
        for p in hdr1.paragraphs:
            p.text=''
        p = hdr1.paragraphs[0] if hdr1.paragraphs else hdr1.add_paragraph()
        p.text='Header sec1 before '
    doc.save(out_path)
    print('saved with headers', out_path)

add_tb_to_header(base, 'C:/Temp/gen2/with_headers.docx')

# Now patch the header xml to insert topAndBottom anchors
def inject_tb_anchors(docx_path, out_path):
    z=zipfile.ZipFile(docx_path,'r')
    parts={n:z.read(n) for n in z.namelist()}
    z.close()
    # List parts
    print('parts', [n for n in parts if 'header' in n or 'footer' in n or 'document' in n])
    for name in [n for n in parts if 'header' in n or 'footer' in n]:
        print(name, parts[name][:1000].decode(errors='ignore')[:1000])
    # For each header, inject anchors
    import copy
    for hdr_name in [n for n in parts if n.startswith('word/header') and n.endswith('.xml')]:
        xml = parts[hdr_name].decode()
        root = ET.fromstring(xml)
        NS_W='http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        # Find first p
        p = root.find(f'{{{NS_W}}}p')
        if p is None:
            continue
        # We'll add three runs with anchors: covering relativeFrom page left, margin center, paragraph offset
        # To keep valid, we need to add relationships for each image
        # Instead of manually adding rels, we will reuse existing media if exists? We'll add new media files and update rels
        # Create new media files in parts: word/media/tb1.png etc
        # For simplicity, copy tb.png bytes to each media
        # Determine hdr rels name
        rels_name = hdr_name.replace('word/', 'word/_rels/') + '.rels'
        rels_xml = parts.get(rels_name, b'<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>').decode()
        rels_root = ET.fromstring(rels_xml)
        # Helper to add image
        def add_image_anchor(p_elem, rels_root, image_path, rid, relH, relV, alignH=None, alignV=None, offH=None, offV=None, wrap='wrapTopAndBottom', dist=None, cx=1080000, cy=720000):
            data = open(image_path,'rb').read()
            media_name = f'media/{rid}.png'
            parts[f'word/{media_name}'] = data
            # add rel
            rel = ET.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
            rel.set('Id', rid); rel.set('Type','http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'); rel.set('Target', media_name)
            rels_root.append(rel)
            # create r/drawing
            r = ET.Element(f'{{{NS_W}}}r')
            drawing = ET.Element('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
            for k,v in [('distT','0'),('distB','0'),('distL','0'),('distR','0'),('simplePos','0'),('relativeHeight','251658240'),('behindDoc','0'),('locked','0'),('layoutInCell','1'),('allowOverlap','1')]:
                drawing.set(k,v)
            if dist:
                for dk,dv in dist.items():
                    drawing.set(dk, str(dv))
            ext = ET.SubElement(drawing, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'); ext.set('cx',str(cx)); ext.set('cy',str(cy))
            docPr = ET.SubElement(drawing, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr'); docPr.set('id','1'); docPr.set('name','Picture 1')
            cNvGF = ET.SubElement(drawing, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}cNvGraphicFramePr')
            gfl = ET.SubElement(cNvGF, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicFrameLocks'); gfl.set('{http://schemas.openxmlformats.org/drawingml/2006/main}noChangeAspect','1')
            graphic = ET.SubElement(drawing, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphic')
            gd = ET.SubElement(graphic, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData'); gd.set('uri','http://schemas.openxmlformats.org/drawingml/2006/picture')
            pic = ET.SubElement(gd, '{http://schemas.openxmlformats.org/drawingml/2006/picture}pic')
            nv = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}nvPicPr')
            cNvPr = ET.SubElement(nv, '{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPr'); cNvPr.set('id','0'); cNvPr.set('name','tb.png')
            ET.SubElement(nv, '{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPicPr')
            blipFill = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}blipFill')
            blip = ET.SubElement(blipFill, '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'); blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', rid)
            stretch = ET.SubElement(blipFill, '{http://schemas.openxmlformats.org/drawingml/2006/main}stretch'); ET.SubElement(stretch, '{http://schemas.openxmlformats.org/drawingml/2006/main}fillRect')
            spPr = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr')
            xfrm = ET.SubElement(spPr, '{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm'); off = ET.SubElement(xfrm, '{http://schemas.openxmlformats.org/drawingml/2006/main}off'); off.set('x','0'); off.set('y','0'); ext2 = ET.SubElement(xfrm, '{http://schemas.openxmlformats.org/drawingml/2006/main}ext'); ext2.set('cx',str(cx)); ext2.set('cy',str(cy))
            ET.SubElement(spPr, '{http://schemas.openxmlformats.org/drawingml/2006/main}prstGeom').set('prst','rect')
            ET.SubElement(drawing, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos').set('x','0'); drawing.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos').set('y','0')
            posH = ET.SubElement(drawing, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH'); posH.set('relativeFrom', relH)
            if alignH:
                al = ET.SubElement(posH, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align'); al.text=alignH
            else:
                po = ET.SubElement(posH, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset'); po.text=str(offH if offH is not None else 0)
            posV = ET.SubElement(drawing, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV'); posV.set('relativeFrom', relV)
            if alignV:
                al2 = ET.SubElement(posV, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align'); al2.text=alignV
            else:
                po2 = ET.SubElement(posV, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset'); po2.text=str(offV if offV is not None else 0)
            wrap = ET.SubElement(drawing, f'{{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}}{wrap}')
            # add polygon if needed
            r.append(drawing)
            # need to append r to p? Actually drawing is inside r via w:drawing element
            # Our structure: r -> drawing (which is wp:anchor wrapped in w:drawing?) Need w:drawing wrapper
            # Correct: r contains w:drawing which contains wp:anchor
            # We created r with anchor directly, need wrapper
            # Fix: create w:drawing element containing anchor
            # We'll recreate properly
            return r, rels_root

        # Clear existing approach is messy. Let's instead construct properly per anchor using w:drawing wrapper
        # Simplify: wipe and rebuild header with 3 anchors correctly

        # Remove all children and rebuild
        for c in list(root):
            root.remove(c)
        rels_root_new = ET.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationships')
        rels_root_new.set('xmlns','http://schemas.openxmlformats.org/package/2006/relationships')
        # copy existing rels? start fresh, we'll add all
        # We'll create 3 anchors
        specs = [
            ('rId1', 'margin', None, 'center', None, 0, 0, {'distT':'100000','distB':'100000','distL':'100000','distR':'100000'}), # margin center topAndBottom
            ('rId2', 'page', 500000, None, 'page', 200000, None, {'distT':'50000','distB':'50000'}), # page offset
            ('rId3', 'paragraph', 300000, None, 'paragraph', 150000, None, None), # paragraph offset
        ]
        p_new = ET.Element(f'{{{NS_W}}}p')
        pPr = ET.SubElement(p_new, f'{{{NS_W}}}pPr'); ET.SubElement(pPr, f'{{{NS_W}}}pStyle').set(f'{{{NS_W}}}val','Header')
        for rid, relH, offH, alignH, relV, offV, alignV, dist in specs:
            data = open('C:/Temp/gen2/tb.png','rb').read()
            parts[f'word/media/{rid}.png'] = data
            rel = ET.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
            rel.set('Id', rid); rel.set('Type','http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'); rel.set('Target', f'media/{rid}.png')
            rels_root_new.append(rel)
            r = ET.Element(f'{{{NS_W}}}r')
            drawing = ET.Element(f'{{{NS_W}}}drawing')
            anchor = ET.Element('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
            for k,v in [('distT','0'),('distB','0'),('distL','0'),('distR','0'),('simplePos','0'),('relativeHeight','251658240'),('behindDoc','0'),('locked','0'),('layoutInCell','1'),('allowOverlap','1')]:
                anchor.set(k,v)
            if dist:
                for dk,dv in dist.items():
                    anchor.set(dk,dv)
            ext = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'); ext.set('cx','1200000'); ext.set('cy','600000')
            docPr = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr'); docPr.set('id','1'); docPr.set('name','TB')
            cNvGF = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}cNvGraphicFramePr')
            gfl = ET.SubElement(cNvGF, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicFrameLocks'); gfl.set('{http://schemas.openxmlformats.org/drawingml/2006/main}noChangeAspect','1')
            graphic = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphic')
            gd = ET.SubElement(graphic, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData'); gd.set('uri','http://schemas.openxmlformats.org/drawingml/2006/picture')
            pic = ET.SubElement(gd, '{http://schemas.openxmlformats.org/drawingml/2006/picture}pic')
            nv = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}nvPicPr')
            cNvPr = ET.SubElement(nv, '{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPr'); cNvPr.set('id','0'); cNvPr.set('name','tb.png')
            ET.SubElement(nv, '{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPicPr')
            blipFill = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}blipFill')
            blip = ET.SubElement(blipFill, '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'); blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', rid)
            stretch = ET.SubElement(blipFill, '{http://schemas.openxmlformats.org/drawingml/2006/main}stretch'); ET.SubElement(stretch, '{http://schemas.openxmlformats.org/drawingml/2006/main}fillRect')
            spPr = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr')
            xfrm = ET.SubElement(spPr, '{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm'); off = ET.SubElement(xfrm, '{http://schemas.openxmlformats.org/drawingml/2006/main}off'); off.set('x','0'); off.set('y','0'); ext2 = ET.SubElement(xfrm, '{http://schemas.openxmlformats.org/drawingml/2006/main}ext'); ext2.set('cx','1200000'); ext2.set('cy','600000')
            ET.SubElement(spPr, '{http://schemas.openxmlformats.org/drawingml/2006/main}prstGeom').set('prst','rect')
            ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos').set('x','0'); anchor.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos').set('y','0')
            posH = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH'); posH.set('relativeFrom', relH)
            if alignH:
                al = ET.SubElement(posH, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align'); al.text=alignH
            else:
                po = ET.SubElement(posH, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset'); po.text=str(offH)
            posV = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV'); posV.set('relativeFrom', relV)
            if alignV:
                al2 = ET.SubElement(posV, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align'); al2.text=alignV
            else:
                po2 = ET.SubElement(posV, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset'); po2.text=str(offV)
            ET.SubElement(anchor, f'{{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}}wrapTopAndBottom')
            drawing.append(anchor)
            r.append(drawing)
            # add text
            t = ET.SubElement(ET.Element(f'{{{NS_W}}}r'), f'{{{NS_W}}}t')
            # Actually add following text run
            r2 = ET.Element(f'{{{NS_W}}}r'); t2 = ET.SubElement(r2, f'{{{NS_W}}}t'); t2.text = f' TB {relH}/{alignH or offH} '
            p_new.append(r); p_new.append(r2)
        root.append(p_new)
        parts[hdr_name] = ET.tostring(root, encoding='utf-8', xml_declaration=True)
        parts[rels_name] = ET.tostring(rels_root_new, encoding='utf-8', xml_declaration=True)

    # Also need footer similarly (create one footer with one TB)
    for ftr_name in [n for n in list(parts.keys()) if n.startswith('word/footer') and n.endswith('.xml')]:
        # ensure footer has at least one TB
        xml = parts[ftr_name].decode()
        root = ET.fromstring(xml)
        NS_W='http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        for c in list(root):
            root.remove(c)
        p_new = ET.Element(f'{{{NS_W}}}p')
        ET.SubElement(ET.SubElement(p_new, f'{{{NS_W}}}pPr'), f'{{{NS_W}}}pStyle').set(f'{{{NS_W}}}val','Footer')
        # add one footer TB with footer distance test
        hdr_name=ftr_name
        rels_name = hdr_name.replace('word/', 'word/_rels/') + '.rels'
        rels_root_new = ET.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationships')
        rels_root_new.set('xmlns','http://schemas.openxmlformats.org/package/2006/relationships')
        rid='rId1'
        data = open('C:/Temp/gen2/tb2.png','rb').read()
        parts[f'word/media/footer_{rid}.png'] = data
        rel = ET.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
        rel.set('Id', rid); rel.set('Type','http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'); rel.set('Target', f'media/footer_{rid}.png')
        rels_root_new.append(rel)
        r = ET.Element(f'{{{NS_W}}}r')
        drawing = ET.Element(f'{{{NS_W}}}drawing')
        anchor = ET.Element('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
        for k,v in [('distT','114300'),('distB','114300'),('distL','114300'),('distR','114300'),('simplePos','0'),('relativeHeight','251658240'),('behindDoc','0'),('locked','0'),('layoutInCell','1'),('allowOverlap','1')]:
            anchor.set(k,v)
        ext = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'); ext.set('cx','1200000'); ext.set('cy','600000')
        docPr = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr'); docPr.set('id','1'); docPr.set('name','TBF')
        cNvGF = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}cNvGraphicFramePr')
        gfl = ET.SubElement(cNvGF, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicFrameLocks'); gfl.set('{http://schemas.openxmlformats.org/drawingml/2006/main}noChangeAspect','1')
        graphic = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphic')
        gd = ET.SubElement(graphic, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData'); gd.set('uri','http://schemas.openxmlformats.org/drawingml/2006/picture')
        pic = ET.SubElement(gd, '{http://schemas.openxmlformats.org/drawingml/2006/picture}pic')
        nv = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}nvPicPr')
        cNvPr = ET.SubElement(nv, '{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPr'); cNvPr.set('id','0'); cNvPr.set('name','tb2.png')
        ET.SubElement(nv, '{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPicPr')
        blipFill = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}blipFill')
        blip = ET.SubElement(blipFill, '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'); blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', rid)
        stretch = ET.SubElement(blipFill, '{http://schemas.openxmlformats.org/drawingml/2006/main}stretch'); ET.SubElement(stretch, '{http://schemas.openxmlformats.org/drawingml/2006/main}fillRect')
        spPr = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr')
        xfrm = ET.SubElement(spPr, '{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm'); off = ET.SubElement(xfrm, '{http://schemas.openxmlformats.org/drawingml/2006/main}off'); off.set('x','0'); off.set('y','0'); ext2 = ET.SubElement(xfrm, '{http://schemas.openxmlformats.org/drawingml/2006/main}ext'); ext2.set('cx','1200000'); ext2.set('cy','600000')
        ET.SubElement(spPr, '{http://schemas.openxmlformats.org/drawingml/2006/main}prstGeom').set('prst','rect')
        ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos').set('x','0'); anchor.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos').set('y','0')
        posH = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH'); posH.set('relativeFrom','margin'); al = ET.SubElement(posH, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align'); al.text='center'
        posV = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV'); posV.set('relativeFrom','page'); po2 = ET.SubElement(posV, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset'); po2.text='500000'
        ET.SubElement(anchor, f'{{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}}wrapTopAndBottom')
        drawing.append(anchor); r.append(drawing); p_new.append(r)
        r2 = ET.Element(f'{{{NS_W}}}r'); t2 = ET.SubElement(r2, f'{{{NS_W}}}t'); t2.text=' Footer TB '
        p_new.append(r2)
        root.append(p_new)
        parts[ftr_name] = ET.tostring(root, encoding='utf-8', xml_declaration=True)
        parts[rels_name] = ET.tostring(rels_root_new, encoding='utf-8', xml_declaration=True)

    # Write body control topAndBottom: add to document.xml first paragraph after Body start
    # Find document.xml body first p after Body start and inject anchor
    doc_xml = parts['word/document.xml'].decode()
    root = ET.fromstring(doc_xml)
    NS_W='http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = root.find(f'{{{NS_W}}}body')
    # Add body TB in first body p (second p)
    ps = body.findall(f'{{{NS_W}}}p')
    if len(ps)>=1:
        p_body = ps[0]
        # add a run with topAndBottom margin center
        r = ET.Element(f'{{{NS_W}}}r')
        drawing = ET.Element(f'{{{NS_W}}}drawing')
        anchor = ET.Element('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
        for k,v in [('distT','114300'),('distB','114300'),('simplePos','0'),('relativeHeight','251658240'),('behindDoc','0'),('locked','0'),('layoutInCell','1'),('allowOverlap','1')]:
            anchor.set(k,v)
        ext = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'); ext.set('cx','1200000'); ext.set('cy','600000')
        docPr = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr'); docPr.set('id','99'); docPr.set('name','BodyTB')
        cNvGF = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}cNvGraphicFramePr')
        gfl = ET.SubElement(cNvGF, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicFrameLocks'); gfl.set('{http://schemas.openxmlformats.org/drawingml/2006/main}noChangeAspect','1')
        graphic = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphic')
        gd = ET.SubElement(graphic, '{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData'); gd.set('uri','http://schemas.openxmlformats.org/drawingml/2006/picture')
        pic = ET.SubElement(gd, '{http://schemas.openxmlformats.org/drawingml/2006/picture}pic')
        nv = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}nvPicPr')
        cNvPr = ET.SubElement(nv, '{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPr'); cNvPr.set('id','0'); cNvPr.set('name','bodytb.png')
        ET.SubElement(nv, '{http://schemas.openxmlformats.org/drawingml/2006/picture}cNvPicPr')
        blipFill = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}blipFill')
        blip = ET.SubElement(blipFill, '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'); blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', 'rIdBodyTB')
        # need to add media and rel for body
        data = open('C:/Temp/gen2/tb.png','rb').read()
        parts['word/media/bodytb.png'] = data
        # update document rels
        doc_rels = ET.fromstring(parts['word/_rels/document.xml.rels'].decode())
        rel = ET.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
        rel.set('Id','rIdBodyTB'); rel.set('Type','http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'); rel.set('Target','media/bodytb.png')
        doc_rels.append(rel)
        parts['word/_rels/document.xml.rels'] = ET.tostring(doc_rels, encoding='utf-8', xml_declaration=True)
        stretch = ET.SubElement(blipFill, '{http://schemas.openxmlformats.org/drawingml/2006/main}stretch'); ET.SubElement(stretch, '{http://schemas.openxmlformats.org/drawingml/2006/main}fillRect')
        spPr = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr')
        xfrm = ET.SubElement(spPr, '{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm'); off = ET.SubElement(xfrm, '{http://schemas.openxmlformats.org/drawingml/2006/main}off'); off.set('x','0'); off.set('y','0'); ext2 = ET.SubElement(xfrm, '{http://schemas.openxmlformats.org/drawingml/2006/main}ext'); ext2.set('cx','1200000'); ext2.set('cy','600000')
        ET.SubElement(spPr, '{http://schemas.openxmlformats.org/drawingml/2006/main}prstGeom').set('prst','rect')
        ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos').set('x','0'); anchor.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}simplePos').set('y','0')
        posH = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH'); posH.set('relativeFrom','margin'); al = ET.SubElement(posH, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}align'); al.text='center'
        posV = ET.SubElement(anchor, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV'); posV.set('relativeFrom','page'); po2 = ET.SubElement(posV, '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset'); po2.text='1000000'
        ET.SubElement(anchor, f'{{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}}wrapTopAndBottom')
        drawing.append(anchor); r.append(drawing); p_body.append(r)
        parts['word/document.xml'] = ET.tostring(root, encoding='utf-8', xml_declaration=True)

    with zipfile.ZipFile(out_path,'w',zipfile.ZIP_DEFLATED) as zout:
        for name,data in parts.items():
            zout.writestr(name,data)
    print('patched tb doc', out_path)

inject_tb_anchors('C:/Temp/gen2/with_headers.docx','C:/Temp/gen2/hf_tb.docx')
