import sys
sys.path.insert(0,'.')
import tempfile, os
from semantic.pipeline import convert_docx
from playwright.sync_api import sync_playwright

tmp='/tmp/para_layout.docx'
res=convert_docx(tmp)
html=res.html
tf=tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
tf.write(html)
tf.close()
with sync_playwright() as pw:
    browser=pw.chromium.launch()
    page=browser.new_page()
    page.goto('file://' + tf.name.replace('\\','/'))
    # Test each paragraph type
    paras=page.locator('p')
    print('p count', paras.count())
    # first para left indent
    first=page.locator('p').nth(0)
    print('first left', first.evaluate('el => getComputedStyle(el).marginLeft'), 'expected 48px or 36pt')
    print('first outer', first.evaluate('el => el.outerHTML')[:200])
    # second right indent
    second=page.locator('p').nth(1)
    print('second right', second.evaluate('el => getComputedStyle(el).marginRight'))
    # third first-line
    third=page.locator('p').nth(2)
    print('third textIndent', third.evaluate('el => getComputedStyle(el).textIndent'))
    # fourth hanging
    fourth=page.locator('p').nth(3)
    print('fourth marginLeft', fourth.evaluate('el => getComputedStyle(el).marginLeft'))
    print('fourth textIndent', fourth.evaluate('el => getComputedStyle(el).textIndent'))
    # spacing
    fifth=page.locator('p').nth(4)
    print('fifth marginTop', fifth.evaluate('el => getComputedStyle(el).marginTop'))
    print('fifth marginBottom', fifth.evaluate('el => getComputedStyle(el).marginBottom'))
    # line spacing
    for i, name in [(5,'single'),(6,'1.5'),(7,'double'),(8,'exact')]:
        el=page.locator('p').nth(i)
        print(f'{name} lineHeight', el.evaluate('el => getComputedStyle(el).lineHeight'))
    # alignment
    centered=page.locator('p').nth(9)
    print('center textAlign', centered.evaluate('el => getComputedStyle(el).textAlign'))
    right=page.locator('p').nth(10)
    print('right textAlign', right.evaluate('el => getComputedStyle(el).textAlign'))
    justified=page.locator('p').nth(11)
    print('justify textAlign', justified.evaluate('el => getComputedStyle(el).textAlign'))
    # no errors
    browser.close()
os.unlink(tf.name)

# Test table cell paragraph layout
from docx import Document
from docx.shared import Pt
doc=Document()
tbl=doc.add_table(rows=1, cols=1)
cell=tbl.cell(0,0)
p=cell.paragraphs[0]
p.text='Cell indented'
p.paragraph_format.left_indent=Pt(18)
p.alignment=1  # center
doc.save('/tmp/table_para.docx')
res2=convert_docx('/tmp/table_para.docx')
tf2=tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
tf2.write(res2.html)
tf2.close()
with sync_playwright() as pw:
    browser=pw.chromium.launch()
    page=browser.new_page()
    page.goto('file://' + tf2.name.replace('\\','/'))
    td_p=page.locator('td p').first
    print('table cell marginLeft', td_p.evaluate('el => getComputedStyle(el).marginLeft'))
    print('table cell textAlign', td_p.evaluate('el => getComputedStyle(el).textAlign'))
    browser.close()
os.unlink(tf2.name)

# Test header paragraph layout
from docx import Document as D2
doc3=D2()
# need header via section
sec=doc3.sections[0]
hdr=sec.header
hp=hdr.paragraphs[0]
hp.text='Header centered'
hp.alignment=1
hp.paragraph_format.space_after=Pt(6)
doc3.save('/tmp/header_para.docx')
res3=convert_docx('/tmp/header_para.docx')
tf3=tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
tf3.write(res3.html)
tf3.close()
with sync_playwright() as pw:
    browser=pw.chromium.launch()
    page=browser.new_page()
    page.goto('file://' + tf3.name.replace('\\','/'))
    hdr_p=page.locator('header p').first
    print('header textAlign', hdr_p.evaluate('el => getComputedStyle(el).textAlign') if hdr_p.count() else 'no header p')
    print('header marginBottom', hdr_p.evaluate('el => getComputedStyle(el).marginBottom') if hdr_p.count() else 'none')
    browser.close()
os.unlink(tf3.name)

# Test hyperlink paragraph layout
import zipfile, shutil
tmp_hyper='/tmp/hyper_para.docx'
shutil.copy('docx-to-html/tests/fixtures/hyperlinks.docx', tmp_hyper)
# inject indent into first para containing hyperlink
import xml.etree.ElementTree as ET
z=zipfile.ZipFile(tmp_hyper,'r')
doc_xml=z.read('word/document.xml').decode()
root=ET.fromstring(doc_xml)
w=root.tag.split('}')[0]+'}'
# find first p with hyperlink, add ind
for p in root.iter(w+'p'):
    for child in p.iter(w+'hyperlink'):
        # found hyperlink para
        pPr=p.find(w+'pPr')
        if pPr is None:
            pPr=ET.Element(w+'pPr')
            p.insert(0,pPr)
        ind=ET.SubElement(pPr, w+'ind')
        ind.set(w+'left','720')
        ind.set(w+'hanging','360')
        break
    break
new_xml=ET.tostring(root, encoding='unicode')
with zipfile.ZipFile(tmp_hyper,'a') as zz:
    zz.writestr('word/document.xml', new_xml)
res4=convert_docx(tmp_hyper)
tf4=tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
tf4.write(res4.html)
tf4.close()
with sync_playwright() as pw:
    browser=pw.chromium.launch()
    page=browser.new_page()
    page.goto('file://' + tf4.name.replace('\\','/'))
    p_hyper=page.locator('p').first
    print('hyper para marginLeft', p_hyper.evaluate('el => getComputedStyle(el).marginLeft'))
    print('hyper para textIndent', p_hyper.evaluate('el => getComputedStyle(el).textIndent'))
    print('hyper link still href?', 'example.com' in p_hyper.inner_html())
    browser.close()
os.unlink(tf4.name)
